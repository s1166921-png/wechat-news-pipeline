#!/usr/bin/env python3
"""
app.py — 跨境电商热点内容创作工坊 (Flask Server)

3 步向导：发现热点 → 生成文章 → 配图导出
"""
import json, math, os, re, time, hashlib, ssl
import urllib.request, urllib.error
from datetime import datetime, timezone, timedelta
from pathlib import Path

from flask import Flask, request, jsonify, send_from_directory, send_file
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from pipeline_core import extractors as _core_extractors
from pipeline_core import importing as _core_importing
from pipeline_core import quality as _core_quality

# ── 路径常量 ──────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
CONFIG_DIR = BASE_DIR / "config"
OUTPUT_DIR = BASE_DIR / "output"
FRONTEND_DIR = BASE_DIR / "frontend"
CACHE_DIR = BASE_DIR / "output"

for d in [OUTPUT_DIR / "articles", OUTPUT_DIR / "images", CACHE_DIR]:
    d.mkdir(parents=True, exist_ok=True)

# ── 加载 .env 配置 ────────────────────────────────────
load_dotenv(BASE_DIR / ".env")

# ── API Keys ──────────────────────────────────────────
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
ARK_API_KEY = os.environ.get("ARK_API_KEY", "")
ARK_IMAGE_MODEL = os.environ.get("ARK_IMAGE_MODEL", "doubao-seedream-4-0-250828")
PORT = int(os.environ.get("PORT", "8888"))
MIN_RAW_REWRITE_CHARS = 300

# ── Flask App ─────────────────────────────────────────
app = Flask(__name__, static_folder=str(FRONTEND_DIR), static_url_path="")

# —— Debug: 记录所有 API 请求的原始 body ——
import logging
logging.basicConfig(filename=str(BASE_DIR / "server.log"), level=logging.INFO,
                    format="%(asctime)s %(levelname)s %(message)s")


def _request_log_preview(payload):
    """Return a safe one-line preview for API request logging."""
    if not isinstance(payload, dict):
        return str(payload)[:300]

    redacted = {}
    sensitive_fields = {"raw_content", "content", "article_content", "rewritten_markdown", "rewritten_html"}
    for key, value in payload.items():
        if key in sensitive_fields and isinstance(value, str):
            redacted[key] = f"[redacted chars={len(value)}]"
        elif isinstance(value, str) and len(value) > 180:
            redacted[key] = value[:180] + f"...[truncated chars={len(value)}]"
        else:
            redacted[key] = value
    return json.dumps(redacted, ensure_ascii=False)[:500]


@app.before_request
def _log_request():
    if request.path.startswith("/api/"):
        payload = request.get_json(silent=True)
        if payload is None:
            payload = request.get_data(as_text=True)[:200]
        logging.info(f"{request.method} {request.path} body={_request_log_preview(payload)}")

@app.errorhandler(400)
def _bad_request(e):
    """返回 JSON 而非 HTML，避免前端 JSON.parse 崩溃"""
    return jsonify({"error": "Bad Request: 请求格式错误，请刷新页面后重试", "detail": str(e)}), 400

@app.errorhandler(500)
def _server_error(e):
    logging.error(f"500 on {request.path}: {e}")
    return jsonify({"error": "服务器内部错误，请稍后重试"}), 500

CST = timezone(timedelta(hours=8))


def _today_cst_label(now=None):
    """Return today's date in Chinese format for prompts."""
    return _core_importing.today_cst_label(now=now)


def _urlopen_final_url(url, timeout=10, headers=None):
    """Open a URL and return the final URL after redirects."""
    return _core_importing.urlopen_final_url(url, timeout=timeout, headers=headers)


def _resolve_wechat_url(url, timeout=10):
    """Normalize WeChat article URLs, including Sogou WeChat redirect links."""
    return _core_importing.resolve_wechat_url(
        url,
        timeout=timeout,
        urlopen_func=_urlopen_final_url,
    )


def _classify_article_import(url="", raw_content=""):
    """Describe how an article should enter the rewrite pipeline."""
    return _core_importing.classify_article_import(url=url, raw_content=raw_content)

# ╔══════════════════════════════════════════════════════╗
# ║  LLM Client (DeepSeek)                              ║
# ╚══════════════════════════════════════════════════════╝

def llm_chat_text(system="", user="", model="deepseek-chat", temperature=0.7, max_tokens=3000, api_key=""):
    """调用 DeepSeek Chat API，返回纯文本。"""
    key = api_key or DEEPSEEK_API_KEY
    if not key:
        raise RuntimeError("DEEPSEEK_API_KEY 未设置")

    effective_max = min(max_tokens, 4096)
    payload = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": temperature,
        "max_tokens": effective_max,
    }).encode("utf-8")

    ctx = ssl.create_default_context()
    for attempt in range(1, 4):
        try:
            req = urllib.request.Request(
                "https://api.deepseek.com/v1/chat/completions",
                data=payload,
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {key}"},
            )
            resp = urllib.request.urlopen(req, timeout=180, context=ctx)
            result = json.loads(resp.read().decode("utf-8"))
            content = result["choices"][0]["message"]["content"]
            if content and len(content.strip()) >= 20:
                return content
            print(f"  [LLM] 响应过短 ({len(content) if content else 0} 字符)")
            return content or ""
        except urllib.error.HTTPError as e:
            if e.code == 429:
                wait = 2 ** attempt
                print(f"  [LLM] 429 限流，等待 {wait}s...")
                time.sleep(wait)
            elif e.code >= 500:
                time.sleep(2)
            else:
                raise
        except Exception as e:
            if attempt < 3:
                print(f"  [LLM] 网络错误: {e}，重试 {attempt}/3...")
                time.sleep(1)
            else:
                print(f"  [LLM] 失败: {e}")
                return ""

    return ""


# ╔══════════════════════════════════════════════════════╗
# ║  Search Engine (GEO-grade multi-engine)               ║
# ╚══════════════════════════════════════════════════════╝

_GN_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/125.0.0.0 Safari/537.36",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "no-cache, no-store, must-revalidate",
    "Pragma": "no-cache",
}

_ECOMMERCE_KEYWORDS = {"跨境电商", "跨境", "亚马逊", "Amazon", "FBA", "Shopee", "eBay",
                       "外贸", "出海", "海外仓", "一件代发", "TikTok", "SHEIN", "Temu",
                       "出口退税", "出口", "退税", "海关", "报关", "VAT", "vat",
                       "关税", "合规", "物流", "选品", "独立站", "品牌出海", "供应链"}

# Smart query expansion: keywords that need industry context to avoid
# being interpreted as generic dictionary words by search engines
_ECOMMERCE_CONTEXT = {
    "选品": "跨境电商选品 产品调研",
    "物流": "跨境物流 国际货运",
    "仓储": "跨境电商仓储 海外仓",
    "财税": "跨境电商财税 合规",
    "合规": "跨境电商合规 法规",
    "VAT": "跨境电商 VAT 欧洲税务",
    "vat": "跨境电商 VAT 欧洲税务",
    "关税": "跨境电商关税 贸易政策",
    "独立站": "跨境电商独立站 建站",
    "品牌出海": "中国品牌出海 跨境电商 DTC",
    "供应链": "跨境供应链 物流管理",
    "运营": "跨境电商运营 亚马逊",
    "TikTok": "TikTok电商 跨境",
    "tiktok": "TikTok电商 跨境",
}

_ROTATING_QUERIES = [
    "跨境电商 最新 动态", "亚马逊 政策 变化", "海外仓 物流 趋势",
    "TikTok Shop 电商", "品牌出海 中国", "独立站 运营 增长",
    "跨境支付 合规", "供应链 出海", "VAT 税务 欧洲",
    "跨境电商 选品", "关税 贸易 政策", "SHEIN Temu 跨境电商",
    "AI 跨境电商 应用", "跨境营销 社交媒体", "Shopee Lazada 东南亚",
]

_rotation_index = 0


def _simple_get(url, timeout=15):
    """Simple HTTP GET returning response text."""
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    try:
        req = urllib.request.Request(url, headers=_GN_HEADERS)
        resp = urllib.request.urlopen(req, timeout=timeout, context=ctx)
        data = resp.read()
        print(f"  [HTTP] {resp.status} {len(data)}B from {url[:80]}")
        return data.decode("utf-8", errors="replace")
    except urllib.error.HTTPError as e:
        print(f"  [HTTP] HTTP {e.code} from {url[:80]}: {e.reason}")
        return ""
    except Exception as e:
        print(f"  [HTTP] {type(e).__name__}: {e} from {url[:80]}")
        return ""


def _search_google_news(keyword, max_results=25):
    """Search Google News RSS for Chinese news articles.

    Google News RSS returns 50-100 recent Chinese news articles reliably.
    """
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS

    results = []
    try:
        # Search with hl=zh-CN for Chinese results but WITHOUT gl=CN
        # (gl=CN severely limits the news pool because of GFW censorship)
        # Add timestamp to bust Google CDN cache (real-time freshness)
        _cache_buster = int(time.time())
        url = f"https://news.google.com/rss/search?q={_quote(keyword)}&hl=zh-CN&ceid=CN:zh-Hans&t={_cache_buster}"
        r = _simple_get(url, timeout=5)
        if not r:
            return results

        soup = _BS(r, "xml")
        for item in soup.select("item"):
            if len(results) >= max_results:
                break
            title_el = item.select_one("title")
            link_el = item.select_one("link")
            source_el = item.select_one("source")
            date_el = item.select_one("pubDate")
            desc_el = item.select_one("description")

            title = title_el.get_text(strip=True) if title_el else ""
            link = link_el.get_text(strip=True) if link_el else ""
            source = source_el.get_text(strip=True) if source_el else ""
            date_str = date_el.get_text(strip=True) if date_el else ""
            # Extract snippet from RSS <description> (HTML, strip tags for plain text)
            desc_html = desc_el.get_text(strip=True) if desc_el else ""
            snippet = ""
            if desc_html:
                # RSS description is HTML — strip all tags, keep text
                from bs4 import BeautifulSoup as _BS2
                try:
                    _desc_soup = _BS2(desc_html, "lxml")
                    snippet = _desc_soup.get_text(" ", strip=True)[:500]
                except Exception:
                    # Fallback: regex strip tags
                    snippet = _re.sub(r"<[^>]+>", " ", desc_html)[:500]
                    snippet = _re.sub(r"\s+", " ", snippet).strip()

            if not title or len(title) < 8:
                continue
            if not link:
                continue
            if "Google News" in title or title.startswith("["):
                continue

            source = source.split(" - ")[0].strip()
            results.append({
                "title": title, "url": link, "source": source,
                "date": date_str, "source_type": "google_news",
                "snippet": snippet,
            })
    except Exception as e:
        print(f"  [GN] Error: {e}")
    return results


def _search_bing(keyword, max_results=10):
    """Search Bing for Chinese results as fallback."""
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS

    results = []
    try:
        url = f"https://cn.bing.com/search?q={_quote(keyword)}&count={max_results}&setlang=zh-cn"
        html = _simple_get(url, timeout=15)
        if not html:
            return results

        soup = _BS(html, "lxml")
        items = []
        items.extend(soup.select("li.b_algo"))
        items.extend(soup.select("ol#b_results > li.b_algo"))
        if not items:
            items = soup.select("h2 a, h3 a")

        for item in items[:max_results]:
            a = item.select_one("h2 a") or item.select_one("h3 a") or item.select_one("a")
            if not a:
                continue
            title = a.get_text(strip=True)
            href = a.get("href", "")
            if not title or len(title) < 8 or not href or not href.startswith("http"):
                continue
            if any(b in href for b in ("bing.com", "microsoft.com/bing", "duckduckgo.com")):
                continue

            cite_el = item.select_one(".b_attribution cite, cite")
            source = cite_el.get_text(strip=True) if cite_el else ""
            # Extract snippet: Bing uses .b_caption p, .b_snippet, or .b_lineclamp2
            snippet_el = item.select_one(".b_caption p, .b_snippet, .b_lineclamp2, p[class*=snippet]")
            snippet = snippet_el.get_text(" ", strip=True)[:500] if snippet_el else ""

            results.append({
                "title": title, "url": href, "source": source,
                "date": "", "source_type": "bing",
                "snippet": snippet,
            })
    except Exception as e:
        print(f"  [Bing] Error: {e}")
    return results


def _search_360(keyword, max_results=10):
    """Search 360 News (news.so.com) for dated Chinese news articles.

    360 的 HTML 结构:
      <a title="完整标题">
        <h3><div class="g-txt-inner g-ellipsis">标题文本</div></h3>
        <p class="summary g-ellipsis3">正文预览摘要</p>
      </a>
    其中 title 属性就是干净完整的文章标题，不需要从混杂的 get_text() 中解析。
    """
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS
    import re as _re

    results = []
    try:
        # Add random param to bust 360 cache
        _cb = _re.sub(r'\.', '', str(time.time()))[-8:]
        url = f"https://news.so.com/ns?q={_quote(keyword)}&_t={_cb}"
        html = _simple_get(url, timeout=10)
        if not html:
            return results

        soup = _BS(html, "lxml")

        # Primary: links inside div.result_wrap (the main news list)
        result_wrap = soup.select_one("div.result_wrap")
        if result_wrap:
            items = result_wrap.select("a[href^=http]")
        else:
            # Fallback: li.res-list
            items = [li.select_one("a") for li in soup.select("li.res-list")]
            items = [a for a in items if a is not None]

        for a in items[:max_results]:
            href = a.get("href", "")
            if not href or not href.startswith("http"):
                continue

            # ★ 优先从 title 属性获取完整标题（360 把干净标题放在这里）
            title = (a.get("title") or "").strip()
            # 备选：从 .g-txt-inner 元素中提取
            if len(title) < 8:
                title_el = a.select_one(".g-txt-inner")
                if title_el:
                    title = title_el.get_text(strip=True)

            if len(title) < 8:
                continue  # 标题太短，跳过

            # 提取正文摘要 (snippet) — 用于 AI 解读
            summary_el = a.select_one(".summary, p[class*=summary]")
            if summary_el:
                snippet = summary_el.get_text(strip=True)
            else:
                snippet = a.get_text(strip=True)

            # 从 link text 末尾解析来源和时间
            link_text = a.get_text(strip=True)
            source = ""
            date_str = ""
            time_match = _re.search(
                r'([一-鿿\w]{2,12})?'
                r'(\d+分钟前|\d+小时前|\d+天前|'
                r'\d{4}-\d{2}-\d{2}\s*\d{2}:\d{2}|'
                r'\d{4}年\d{1,2}月\d{1,2}日)$',
                link_text
            )
            if time_match:
                source = (time_match.group(1) or "").strip()
                date_str = time_match.group(2)
                if source and len(source) <= 4 and not _re.search(r'[一-鿿]', source):
                    source = ""

            results.append({
                "title": title.strip(),
                "url": href,
                "source": source,
                "date": date_str,
                "source_type": "360search",
                "snippet": snippet,
            })
    except Exception as e:
        print(f"  [360] Error: {e}")
    return results


def _search_sogou_news(keyword, max_results=10):
    """Search Sogou News (news.sogou.com) for dated Chinese news articles."""
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS
    import re as _re

    results = []
    try:
        url = f"https://news.sogou.com/news?query={_quote(keyword)}&mode=1"
        html = _simple_get(url, timeout=10)
        if not html:
            return results

        soup = _BS(html, "lxml")

        # Primary: links inside div.wrap (the main news list area)
        wrap = soup.select_one("div.wrap")
        if wrap:
            items = wrap.select("a[href^=http]")
        else:
            # Fallbacks
            items = soup.select("div.news-item a[href^=http], div[class*=news] a[href^=http]")
        if not items:
            items = soup.select("a[href^=http]")

        seen_urls = set()
        for a in items[:max_results * 2]:
            title = a.get_text(strip=True)
            href = a.get("href", "")
            if len(title) < 10 or not href or not href.startswith("http"):
                continue
            if href in seen_urls:
                continue
            seen_urls.add(href)

            # Find date and snippet in ancestor/neighbor elements
            date_str = ""
            source = ""
            snippet = ""
            parent_block = a.parent
            for _ in range(5):
                if not parent_block:
                    break
                ptext = parent_block.get_text(" ", strip=True)
                dm = _re.search(
                    r'(\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{2}-\d{2}|'
                    r'\d{1,2}小时前|\d{1,2}天前|\d{1,2}分钟前)',
                    ptext
                )
                if dm:
                    date_str = dm.group(1)
                    # Try to extract source from the same text
                    src_m = _re.search(r'([一-鿿]{2,10})(?:' + _re.escape(date_str) + r')', ptext)
                    if src_m:
                        source = src_m.group(1)
                    # Extract snippet: text after title, before date/source
                    if title in ptext:
                        after_title = ptext[ptext.index(title) + len(title):]
                        # Remove source and date from end
                        snip = after_title
                        if date_str and date_str in snip:
                            snip = snip[:snip.index(date_str)]
                        if source and source in snip:
                            snip = snip[:snip.rindex(source)]
                        snip = snip.strip(" -｜|,.，。;；:：\n\t ")
                        if len(snip) > 20 and len(snip) < 800:
                            snippet = snip[:500]
                    break
                parent_block = parent_block.parent

            # Fallback: look for desc/summary elements near the link
            if not snippet:
                for selector in [".news-desc", ".des", ".summary", "p[class*=desc]", "p[class*=summary]", "div[class*=abstract]"]:
                    _el = a.parent.select_one(selector) if a.parent else None
                    if _el:
                        snippet = _el.get_text(" ", strip=True)[:500]
                        break

            results.append({
                "title": title[:200],
                "url": href,
                "source": source,
                "date": date_str,
                "source_type": "sogou_news",
                "snippet": snippet,
            })
            if len(results) >= max_results:
                break

    except Exception as e:
        print(f"  [Sogou] Error: {e}")
    return results


def _search_ebrun(keyword, max_results=10):
    """Search 亿邦动力 (m.ebrun.com) for cross-border e-commerce news articles.

    亿邦动力是中国电商行业头部媒体，专注跨境电商/零售/品牌出海。
    Mobile 版本可直接抓取（桌面版有 403 反爬）。
    """
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS
    import re as _re

    results = []
    try:
        url = f"https://m.ebrun.com/search?keyword={_quote(keyword)}"
        html = _simple_get(url, timeout=12)
        if not html:
            return results

        soup = _BS(html, "lxml")
        # Article links have substantial titles (>15 chars) and href to ebrun.com
        article_links = []
        for a in soup.select("a[href]"):
            title = a.get_text(strip=True)
            href = a.get("href", "")
            if len(title) > 15 and "ebrun.com/" in href and "search" not in href[:30]:
                # Resolve relative URLs
                if href.startswith("/"):
                    href = "https://m.ebrun.com" + href
                article_links.append((title, href))

        # Build CJK bigrams for relevance filtering
        cjk_chars = [c for c in keyword if '一' <= c <= '鿿']
        kw_bigrams = set()
        for i in range(len(cjk_chars) - 1):
            kw_bigrams.add(cjk_chars[i] + cjk_chars[i + 1])

        # Deduplicate by URL
        seen = set()
        for title, href in article_links:
            url_key = _re.sub(r"[?#].*$", "", href).rstrip("/")
            if url_key in seen or len(results) >= max_results:
                continue

            # Relevance filter: title must contain the keyword phrase or enough bigrams
            if len(cjk_chars) >= 3:
                cjk_phrase = ''.join(cjk_chars)
                if cjk_phrase not in title:
                    bigram_hits = sum(1 for bg in kw_bigrams if bg in title)
                    if bigram_hits < max(1, len(kw_bigrams) // 2):
                        continue  # not relevant enough
            elif len(cjk_chars) >= 1:
                if not any(c in title for c in cjk_chars):
                    continue

            # ── Filter out daily digest/roundup articles ──
            title_lower = title.lower()
            _digest_patterns = [
                r'\d{1,2}点\w{1,3}电商',  # "14点电商" "15点聊电商"
                r'\d{1,2}点聊',             # "15点聊"
                r'跨境日报',                # "跨境日报：..."
                r'电商日报',                # "电商日报"
                r'每日\w{1,2}报',           # "每日快报"
                r'早报\s*[|｜]',            # "早报 | ..."
                r'晚报\s*[|｜]',            # "晚报 | ..."
                r'行业周报',                # "行业周报"
                r'一周\w{1,3}闻',           # "一周要闻"
                r'\d{1,2}月\w{1,3}报',      # "5月快报"
            ]
            _is_digest = False
            for _dp in _digest_patterns:
                if _re.search(_dp, title_lower):
                    _is_digest = True
                    break
            if _is_digest:
                continue

            seen.add(url_key)

            # Try to find date in nearby text (search results show "6月26日" etc.)
            date_str = ""
            # Look for date patterns in the page text near this article
            date_match = _re.search(r'(\d{1,2}月\d{1,2}日)', title)
            if not date_match:
                # Search surrounding text in HTML for dates
                idx = html.find(href[:60])
                if idx > 0:
                    nearby = html[max(0, idx - 200):idx + 200]
                    date_match = _re.search(r'(\d{1,2}月\d{1,2}日)', nearby)

            if date_match:
                # Convert "7月20日" → "2026-07-20"
                m, d = date_match.group(1).replace("月", " ").replace("日", "").split()
                date_str = f"2026-{int(m):02d}-{int(d):02d}"

            # ── Extract snippet from nearby description elements ──
            snippet = ""
            idx = html.find(href[:60])
            if idx > 0:
                nearby_html = html[max(0, idx - 100):idx + 600]
                nearby_soup = _BS(nearby_html, "lxml")
                for desc_sel in [".desc", ".description", ".summary", ".abstract", "p[class*=desc]", "p[class*=intro]"]:
                    desc_el = nearby_soup.select_one(desc_sel)
                    if desc_el:
                        snippet = desc_el.get_text(" ", strip=True)[:500]
                        # Don't reuse the title as snippet
                        if snippet and title[:20] in snippet:
                            snippet = snippet.replace(title[:20], "").strip()
                        break
                # Fallback: extract text after the link
                if not snippet:
                    link_el = nearby_soup.select_one(f"a[href*='{href[:40]}']")
                    if link_el:
                        parent_text = link_el.parent.get_text(" ", strip=True) if link_el.parent else ""
                        if parent_text and len(parent_text) > len(title) + 20:
                            snippet = parent_text[len(title):].strip()[:500]

            results.append({
                "title": title[:200],
                "url": href,
                "source": "亿邦动力",
                "date": date_str,
                "source_type": "ebrun",
                "snippet": snippet,
            })

    except Exception as e:
        print(f"  [Ebrun] Error: {e}")
    return results


def _search_36kr(keyword, max_results=10):
    """Search 36Kr RSS feed (36kr.com/feed) for Chinese tech/business news.

    36Kr RSS 返回 30 条最新文章，含 pubDate。按关键词筛选后返回。
    适合作为跨境科技/创投类新闻的补充来源。
    """
    from bs4 import BeautifulSoup as _BS
    import re as _re

    results = []
    try:
        html = _simple_get("https://36kr.com/feed", timeout=10)
        if not html:
            return results

        soup = _BS(html, "xml")
        items = soup.select("item")
        if not items:
            return results

        # Build keyword parts for matching — handles both English (space-split)
        # and Chinese (character bigrams + full phrase) correctly.
        kw_lower = keyword.lower()
        kw_parts = []
        # Space-separated tokens (works for English)
        for part in keyword.split():
            if len(part) >= 2:
                kw_parts.append(part.lower())
        # Chinese: extract CJK bigrams + keep full phrase as primary match
        cjk_chars = [c for c in keyword if '一' <= c <= '鿿' or '㐀' <= c <= '䶿']
        if len(cjk_chars) >= 3:
            # Full CJK phrase is the most important match
            cjk_phrase = ''.join(cjk_chars)
            if cjk_phrase not in kw_parts:
                kw_parts.insert(0, cjk_phrase)
            # Add bigrams for partial matching
            for i in range(len(cjk_chars) - 1):
                bigram = cjk_chars[i] + cjk_chars[i + 1]
                if bigram not in kw_parts:
                    kw_parts.append(bigram)
        # Remove very short parts (single ASCII char etc.)
        kw_parts = [p for p in kw_parts if len(p) >= 2]

        # Require at least the full keyword or a majority of bigrams to match
        MIN_MATCHES = min(2, len(kw_parts)) if kw_parts else 0

        for item in items:
            if len(results) >= max_results:
                break

            title_el = item.select_one("title")
            link_el = item.select_one("link")
            date_el = item.select_one("pubDate")
            desc_el = item.select_one("description")

            title = title_el.get_text(strip=True) if title_el else ""
            link = link_el.get_text(strip=True) if link_el else ""
            date_str = date_el.get_text(strip=True) if date_el else ""
            desc = desc_el.get_text(strip=True) if desc_el else ""

            if len(title) < 10 or not link:
                continue

            # Score relevance: must match the full CJK phrase OR enough bigrams
            combined = (title + " " + desc).lower()
            match_score = sum(1 for p in kw_parts if p in combined) if kw_parts else 0

            # Full CJK phrase match is worth 3x (it's the primary signal)
            if kw_parts and kw_parts[0] in combined:
                match_score += len(kw_parts)  # boost for full phrase match

            # Skip if below minimum match threshold
            if match_score < MIN_MATCHES:
                continue

            # Parse date: "2026-07-20 12:32:43  +0800"
            parsed_date = ""
            date_match = _re.match(r'(\d{4}-\d{2}-\d{2})', date_str)
            if date_match:
                parsed_date = date_match.group(1)

            results.append({
                "title": title[:200],
                "url": link,
                "source": "36氪",
                "date": parsed_date,
                "source_type": "36kr",
                "snippet": desc[:200],
                "_match_score": match_score,
            })

        # Sort by relevance (higher match score first)
        results.sort(key=lambda r: r.get("_match_score", 0), reverse=True)
        for r in results:
            r.pop("_match_score", None)

    except Exception as e:
        print(f"  [36Kr] Error: {e}")
    return results


def _search_wechat(keyword, max_results=10):
    """Search WeChat Official Account articles via Sogou WeChat (weixin.sogou.com).

    Strategy to match WeChat搜一搜 as closely as possible:
    1. QUERY EXPANSION: extract core terms & search multiple queries to get diversity
    2. TIMESTAMP EXTRACTION: parse timeConvert() Unix ts from JS (Sogou hides dates)
    3. MULTI-PAGE FETCH: paginate to find recent articles (Sogou defaults = old)
    4. RECENCY SORT: newest first, with quality boost for known media sources

    HTML structure per result:
      <li id="sogou_vr_...">
        <div class="img-box"> <a data-z="art" href="/link?url=..."><img/></a> </div>
        <div class="txt-box">
          <h3><a href="...">TITLE</a></h3>
          <p class="txt-info">SNIPPET</p>
          <div class="s-p">SOURCE_NAME</div>
        </div>
      </li>
    JavaScript contains: document.write(timeConvert('UNIX_TS')) for each article
    """
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS
    from datetime import datetime as _dt
    import time as _time
    import re as _re

    import urllib.request as _ur
    import ssl as _ssl

    # ── Query expansion ──
    tokens = keyword.strip().split()
    core_queries = []
    if len(tokens) >= 2:
        short_tokens = [t for t in tokens if len(t) <= 8 and not t.isdigit()]
        for st in short_tokens[:2]:
            if st.lower() not in ("的", "了", "是", "在", "和", "与", "及", "或",
                                  "最新", "新闻", "动态", "趋势", "分析", "报告",
                                  "市场", "行业", "发展", "未来", "中国"):
                core_queries.append(st)
    all_queries = [keyword] + [q for q in core_queries if q != keyword]
    all_queries = all_queries[:2]  # max 2 queries

    wx_headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    }
    ctx = _ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = _ssl.CERT_NONE

    _QUALITY_SOURCES = {
        "华尔街见闻", "虎嗅", "虎嗅APP", "36氪", "阑夕", "晚点LatePost",
        "财经杂志", "第一财经", "界面新闻", "澎湃新闻", "新京报", "南方周末",
        "经济观察报", "中国经营报", "每日经济新闻", "21世纪经济报道",
        "电商之家", "电商在线", "亿邦动力", "天下网商", "创业邦",
        "极客公园", "爱范儿", "品玩", "腾讯科技", "新浪科技", "网易科技",
        "钛媒体", "雪球", "东方财富网", "同花顺", "和讯网",
        "海关发布", "商务部", "中国跨境电商综合试验区",
    }

    _NOW_TS = int(_time.time())

    def _parse_page(html):
        """Parse one page of Sogou WeChat HTML. Returns (results, timestamps)."""
        out = []
        soup = _BS(html, "lxml")
        items = soup.select("ul.news-list li")
        if not items:
            items = soup.select("div.news-box li, div.txt-box li")

        # Extract all timestamps from JavaScript timeConvert() calls
        timestamps = []
        for m in _re.finditer(r"timeConvert\('(\d+)'\)", html):
            try:
                timestamps.append(int(m.group(1)))
            except ValueError:
                timestamps.append(0)

        for i, li in enumerate(items):
            h3 = li.select_one("h3")
            title = h3.get_text(strip=True) if h3 else ""
            href = ""
            if h3:
                a = h3.select_one("a") or h3.find_parent("a")
                if not a:
                    a = li.select_one("a[href]")
                if a:
                    href = a.get("href", "")

            if len(title) < 8:
                best_len = 0
                for a in li.select("a[href]"):
                    a_text = a.get_text(strip=True)
                    if len(a_text) > best_len:
                        best_len = len(a_text)
                        title = a_text
                        href = a.get("href", "")

            if len(title) < 8:
                continue

            if href.startswith("/"):
                href = "https://weixin.sogou.com" + href

            sp = li.select_one("div.s-p")
            source = sp.get_text(strip=True) if sp else ""
            if not source:
                full_text = li.get_text(strip=True)
                src_match = _re.search(r'([一-鿿\w]{2,20})\s*$', full_text)
                if src_match:
                    candidate = src_match.group(1)
                    if not any(kw in candidate for kw in ["阅读", "点赞", "在看", "全文", "更多", "相关"]):
                        source = candidate

            txt_info = li.select_one("p.txt-info")
            snippet = txt_info.get_text(strip=True) if txt_info else ""
            if not snippet:
                full_text = li.get_text(strip=True)
                snippet = full_text
                if title in snippet:
                    snippet = snippet[snippet.index(title) + len(title):]
                if source and source in snippet:
                    snippet = snippet[:snippet.rindex(source)].strip()
            snippet = snippet[:200]

            # Timestamp: i-th item → i-th timeConvert() call
            ts = timestamps[i] if i < len(timestamps) else 0
            age_days = (_NOW_TS - ts) // 86400 if ts else 9999
            date_str = _dt.fromtimestamp(ts).strftime("%Y-%m-%d") if ts else ""

            is_quality = any(qs in source for qs in _QUALITY_SOURCES)

            out.append({
                "title": title.strip(),
                "url": href,
                "source": source if source else "微信公众号",
                "date": date_str,
                "source_type": "wechat",
                "snippet": snippet,
                "_ts": ts,
                "_age_days": age_days,
                "_quality": is_quality,
            })
        return out

    # ── Multi-query + multi-page search ──
    all_results = []
    seen_titles = set()
    pages_fetched = 0
    MAX_PAGES = 6

    for qi, query in enumerate(all_queries):
        for page in range(1, 4):  # up to 3 pages per query
            if pages_fetched >= MAX_PAGES:
                break
            try:
                url = f"https://weixin.sogou.com/weixin?type=2&query={_quote(query)}&page={page}&ie=utf8"
                req = _ur.Request(url, headers=wx_headers)
                resp = _ur.urlopen(req, timeout=10, context=ctx)
                html = resp.read().decode("utf-8", errors="replace")
                pages_fetched += 1

                if not html or "请输入验证码" in html or "antibot" in html.lower():
                    break  # Stop pagination for this query

                parsed = _parse_page(html)

                # If page returned 0 items, stop pagination
                if not parsed:
                    break

                # Check if this page has any recent (< 30 days) results
                recent_count = sum(1 for r in parsed if r["_age_days"] < 30)

                for r in parsed:
                    title_key = _re.sub(r'[^一-鿿\w]', '', r["title"])[:20]
                    if title_key not in seen_titles:
                        seen_titles.add(title_key)
                        all_results.append(r)

                # Stop pagination if this page had recent results (we got what we need)
                # Or if all results on this page are > 365 days old
                if recent_count > 0 or all(r["_age_days"] > 365 for r in parsed):
                    break

            except Exception as e:
                if page == 1:
                    print(f"  [WeChat] Query '{query[:30]}' error: {e}")
                break  # Stop pagination on error

    # ── CJK bigram relevance filter ──
    # Sogou WeChat returns broad results; filter to those with keyword match
    _kw_cjk = [c for c in keyword if '一' <= c <= '鿿']
    _kw_bigrams = set()
    for i in range(len(_kw_cjk) - 1):
        _kw_bigrams.add(_kw_cjk[i] + _kw_cjk[i + 1])
    _kw_lower = keyword.lower()
    _kw_english = [t for t in keyword.split() if t.isascii() and len(t) >= 2]

    _filtered = []
    for r in all_results:
        title = r["title"]
        snippet = r.get("snippet", "")
        combined = title + " " + snippet

        # English keyword matching (e.g., "temu", "FBA")
        # Require ALL English terms to match — they're highly specific and discriminative
        _eng_match = True
        if _kw_english:
            _eng_match = all(t.lower() in combined.lower() for t in _kw_english)

        # CJK matching — must satisfy BOTH CJK and English matches
        _cjk_match = True
        if len(_kw_cjk) >= 3:
            _cjk_phrase = ''.join(_kw_cjk)
            if _cjk_phrase in combined:
                pass  # full phrase match
            else:
                bigram_hits = sum(1 for bg in _kw_bigrams if bg in combined)
                min_bigrams = max(1, len(_kw_bigrams) // 2)
                if bigram_hits < min_bigrams:
                    _cjk_match = False
        elif len(_kw_cjk) >= 1:
            if not any(c in combined for c in _kw_cjk):
                _cjk_match = False

        if not _cjk_match or not _eng_match:
            continue

        _filtered.append(r)

    if _kw_cjk or _kw_english:
        before_filter = len(all_results)
        all_results = _filtered
        if before_filter > len(all_results):
            print(f"  [WeChat] CJK filter: {before_filter} → {len(all_results)} results")

    # ── Sort: recency first, quality boost within same time bucket ──
    def _sort_key(r):
        ts = r.get("_ts", 0)
        quality = 1 if r.get("_quality") else 0
        # Round to day buckets so quality sorts within same day
        day_bucket = ts // 86400 if ts else 0
        return (day_bucket, quality)

    all_results.sort(key=_sort_key, reverse=True)

    # Count recent results
    recent = sum(1 for r in all_results if r.get("_age_days", 9999) < 90)
    today_count = sum(1 for r in all_results if r.get("_age_days", 9999) < 7)

    # Clean internal fields
    for r in all_results:
        r.pop("_ts", None)
        r.pop("_age_days", None)
        r.pop("_quality", None)

    print(f"  [WeChat] {len(all_queries)}q×{pages_fetched}p → {len(all_results)} unique "
          f"({recent} <90d, {today_count} <7d) for '{keyword[:40]}'")
    return all_results[:max_results]


# ═══════════════════════════════════════════════════════════
# 第 1 层：多搜索引擎交叉覆盖 — site:mp.weixin.qq.com
# ═══════════════════════════════════════════════════════════
# 每个搜索引擎独立爬取 mp.weixin.qq.com，覆盖面互不重叠。
# 搜狗微信是腾讯直供给搜狗的子集，Bing/Google/360 则通过各自
# 的爬虫从全网链接中独立发现微信文章——它们的索引可能包含搜狗没有的 URL。

def _search_bing_wechat(keyword, max_results=5):
    """Bing site:mp.weixin.qq.com — 搜 Bing 索引里的微信公众号文章。

    Bing 通过全球爬虫链路发现 mp.weixin.qq.com 页面，覆盖面与搜狗无直接关联。
    """
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS
    import re as _re

    results = []
    # 用 site: 操作符限定只搜微信公众号域名
    query = f'site:mp.weixin.qq.com {keyword}'
    try:
        url = f"https://cn.bing.com/search?q={_quote(query)}&count={max_results + 5}&setlang=zh-cn"
        html = _simple_get(url, timeout=12)
        if not html:
            return results

        soup = _BS(html, "lxml")
        items = soup.select("li.b_algo, ol#b_results > li.b_algo")
        if not items:
            items = [a.find_parent("li") or a for a in soup.select("h2 a[href*='mp.weixin.qq.com']")]

        for item in items[:max_results]:
            a = item.select_one("h2 a") or item.select_one("a")
            if not a:
                continue
            title = a.get_text(strip=True)
            href = a.get("href", "")
            if not title or len(title) < 6:
                continue
            # 只保留微信文章链接
            if "mp.weixin.qq.com" not in href:
                continue

            # 从 citation 中提取来源公众号名
            cite_el = item.select_one("cite, .b_attribution cite, .b_caption cite")
            source = cite_el.get_text(strip=True) if cite_el else ""

            # 尝试从 source 中提取公众号名（格式: mp.weixin.qq.com · 公众号名）
            if "·" in source:
                source = source.split("·")[-1].strip()
            elif "mp.weixin" in source:
                source = ""

            # snippet
            snippet_el = item.select_one(".b_caption p, .b_snippet, .b_lineclamp2")
            snippet = snippet_el.get_text(" ", strip=True)[:500] if snippet_el else ""

            results.append({
                "title": title, "url": href, "source": source,
                "date": "", "source_type": "bing_wechat",
                "snippet": snippet,
            })
    except Exception as e:
        print(f"  [Bing-WeChat] Error: {e}")
    return results


def _search_google_wechat(keyword, max_results=5):
    """Google site:mp.weixin.qq.com — 搜 Google 索引里的微信公众号文章。

    Google 拥有最大的网页索引库，可能收录了搜狗和 Bing 都没发现的文章。
    使用非 JS 的 Google 搜索（通过 html 抓取），避免 Selenium 依赖。
    """
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS
    import re as _re

    results = []
    query = f'site:mp.weixin.qq.com {keyword}'
    try:
        # 使用 Google 的纯 HTML 搜索接口（不依赖 JS）
        url = f"https://www.google.com/search?q={_quote(query)}&num={max_results + 5}&hl=zh-CN"
        html = _simple_get(url, timeout=12)
        if not html:
            return results

        soup = _BS(html, "lxml")

        # Google 搜索结果的主要容器
        items = []
        # 主结果区
        for container in soup.select("div.g, div[data-sokoban-container], div.MjjYud"):
            # 找链接
            a = container.select_one("a[href*='mp.weixin.qq.com']")
            if a and a not in [i.select_one("a") for i in items if i.select_one("a")]:
                items.append(container)

        if not items:
            # fallback: 直接找所有链接
            for a in soup.select("a[href*='mp.weixin.qq.com/s/']"):
                parent = a.find_parent("div")
                while parent and parent not in items:
                    if parent.name in ("div", "li", "section"):
                        items.append(parent)
                        break
                    parent = parent.find_parent()

        for item in items[:max_results]:
            a = item.select_one("a[href*='mp.weixin.qq.com/s/']")
            if not a:
                continue
            title = a.get_text(strip=True)
            href = a.get("href", "")
            if not title or len(title) < 6:
                continue
            # 清理 Google 的 /url?q= 包装
            if "/url?q=" in href:
                m = _re.search(r'/url\?q=([^&]+)', href)
                href = m.group(1) if m else href

            if "mp.weixin.qq.com" not in href:
                continue

            # 尝试提取公众号名或来源
            source = ""
            cite_el = item.select_one("cite, .iUh30, span[role='text']")
            if cite_el:
                cite_text = cite_el.get_text(strip=True)
                # 格式通常是 mp.weixin.qq.com › ...
                if "›" in cite_text:
                    parts = cite_text.split("›")
                    if len(parts) >= 2:
                        source = parts[1].strip()

            snippet_el = item.select_one("span.aCOpRe, div.VwiC3b, span.st, div[data-sncf]")
            snippet = snippet_el.get_text(" ", strip=True)[:500] if snippet_el else ""

            results.append({
                "title": title, "url": href, "source": source,
                "date": "", "source_type": "google_wechat",
                "snippet": snippet,
            })
    except Exception as e:
        print(f"  [Google-WeChat] Error: {e}")
    return results


def _search_360_wechat(keyword, max_results=5):
    """360 site:mp.weixin.qq.com — 搜 360 索引里的微信公众号文章。

    360 搜索在中国市场占有率高，对中文内容的索引可能与搜狗/Bing 形成互补。
    """
    from urllib.parse import quote as _quote
    from bs4 import BeautifulSoup as _BS
    import re as _re

    results = []
    query = f'site:mp.weixin.qq.com {keyword}'
    try:
        _cb = _re.sub(r'\.', '', str(time.time()))[-8:]
        url = f"https://www.so.com/s?q={_quote(query)}&pn=1&_t={_cb}"
        html = _simple_get(url, timeout=12)
        if not html:
            return results

        soup = _BS(html, "lxml")

        # 360 搜索结果容器
        items = soup.select("li.res-list, div.result, div.results")
        if not items:
            # 降级：找所有指向 mp.weixin.qq.com 的链接
            items = []
            for a in soup.select("a[href*='mp.weixin.qq.com']"):
                parent = a.find_parent(["li", "div.result", "div.res-list"])
                if parent and parent not in items:
                    items.append(parent)

        for item in items[:max_results]:
            a = item.select_one("a[href*='mp.weixin.qq.com']")
            if not a:
                continue
            title = (a.get("title") or a.get_text(strip=True) or "").strip()
            href = a.get("href", "")
            if not title or len(title) < 6:
                continue
            if "mp.weixin.qq.com" not in href:
                continue

            # snippet
            summary_el = item.select_one(".res-desc, .res-summary, .summary, p[class*=desc]")
            snippet = summary_el.get_text(strip=True)[:500] if summary_el else ""

            # 来源
            source = ""
            cite_el = item.select_one("cite, .res-linkinfo, .attr")
            if cite_el:
                source_text = cite_el.get_text(strip=True)
                if "mp.weixin" in source_text:
                    source = ""
                else:
                    source = source_text[:50]

            results.append({
                "title": title, "url": href, "source": source,
                "date": "", "source_type": "360_wechat",
                "snippet": snippet,
            })
    except Exception as e:
        print(f"  [360-WeChat] Error: {e}")
    return results


# ═══════════════════════════════════════════════════════════
# 第 2 层：反向链接发现 — 从社交平台提取微信文章引用
# ═══════════════════════════════════════════════════════════


def _discover_wechat_from_xhs(keyword, max_results=5):
    """从小红书搜索结果中提取引用的微信公众号文章链接。

    策略：
    1. 用 OpenCLI 搜小红书，获取笔记内容
    2. 从笔记正文/描述中提取 mp.weixin.qq.com 链接
    3. 每个链接对应一篇公众号文章——被别人引用过的文章往往质量更高
    """
    import subprocess as _sp
    import json as _json
    import re as _re

    results = []
    try:
        # 用 opencli 搜小红书（--profile 指定了 xg9447p9 即 zhicheng 用户）
        proc = _sp.run(
            ["opencli", "--profile", "xg9447p9", "xiaohongshu", "search", keyword, "-f", "json",
             "--limit", str(min(max_results + 10, 20))],
            capture_output=True, text=True, timeout=45,
            creationflags=_sp.CREATE_NO_WINDOW if hasattr(_sp, "CREATE_NO_WINDOW") else 0,
        )
        if proc.returncode != 0 or not proc.stdout:
            # 降级：try yaml format
            proc = _sp.run(
                ["opencli", "--profile", "xg9447p9", "xiaohongshu", "search", keyword, "-f", "yaml",
                 "--limit", str(min(max_results + 10, 20))],
                capture_output=True, text=True, timeout=45,
                creationflags=_sp.CREATE_NO_WINDOW if hasattr(_sp, "CREATE_NO_WINDOW") else 0,
            )

        raw_output = proc.stdout
        if not raw_output:
            return results

        # 从输出中提取所有 mp.weixin.qq.com/s/... 的 URL
        wx_urls = set()
        for m in _re.finditer(r'https?://mp\.weixin\.qq\.com/s/[^\s"\'\]\)>，。）】&]+', raw_output):
            wx_urls.add(m.group(0))

        # 如果有 JSON 输出，尝试解析结构
        try:
            data = _json.loads(raw_output)
            if isinstance(data, list):
                for note in data:
                    title = note.get("title", "")
                    author = note.get("author", "")
                    note_url = note.get("url", "")
                    likes = note.get("likes", "")
                    # 从 title/description 中提取微信链接
                    desc = note.get("description", "")
                    combined = f"{title} {desc}"
                    for m in _re.finditer(r'https?://mp\.weixin\.qq\.com/s/[^\s"\'\]\)>，。）】&]+', combined):
                        wx_url = m.group(0)
                        if wx_url not in wx_urls:
                            wx_urls.add(wx_url)
                            results.append({
                                "title": f"[小红书引用] {title[:80]}",
                                "url": wx_url,
                                "source": f"XHS/{author}" if author else "XHS",
                                "date": note.get("published_at", ""),
                                "source_type": "xhs_discovery",
                                "snippet": f"小红书 @{author}: {desc[:200]}",
                            })
            return results[:max_results]
        except (_json.JSONDecodeError, Exception):
            pass

        # 非 JSON 输出：直接处理提取到的 URL
        for wx_url in list(wx_urls)[:max_results]:
            results.append({
                "title": f"[小红书引用] {keyword}",
                "url": wx_url,
                "source": "XHS",
                "date": "",
                "source_type": "xhs_discovery",
                "snippet": "",
            })
    except _sp.TimeoutExpired:
        print(f"  [XHS-Discovery] Timeout for '{keyword[:40]}'")
    except FileNotFoundError:
        print(f"  [XHS-Discovery] opencli not found; skip")
    except Exception as e:
        print(f"  [XHS-Discovery] Error: {e}")
    return results


# ═══════════════════════════════════════════════════════════
# 第 3 层：内链爬取 — 从已发现的文章扩散到更多文章
# ═══════════════════════════════════════════════════════════


def _extract_internal_links(article_url, max_links=5):
    """从一篇微信公众号文章中提取它引用的其他公众号文章链接。

    微信公众号文章常见的内部链接模式：
    - 推荐阅读（卡片形式）
    - 相关阅读
    - 文内超链接（引用其他号的文章）
    - 往期回顾
    - 卡片链接（mp.weixin.qq.com/s?...）

    返回: list of {title, url, source, ...}
    """
    import re as _re

    results = []
    try:
        html = _simple_get(article_url, timeout=10)
        if not html:
            return results

        # 提取所有 mp.weixin.qq.com/s/... 链接
        link_pattern = re.compile(
            r'<a[^>]*href="(https?://mp\.weixin\.qq\.com/s/[^"]*)"[^>]*>'
            r'(.*?)'
            r'</a>',
            re.DOTALL,
        )
        seen_urls = set()
        for m in link_pattern.finditer(html):
            url = m.group(1)
            # 规范化 URL：去掉 amp; 等
            url = url.replace("&amp;", "&")
            # 提取链接文字（去 HTML 标签）
            link_text = _re.sub(r'<[^>]+>', '', m.group(2)).strip()

            if url in seen_urls:
                continue
            # 跳过当前文章自身
            if url.rstrip("/") == article_url.rstrip("/"):
                continue
            seen_urls.add(url)

            if link_text and len(link_text) >= 4:
                results.append({
                    "title": link_text[:120],
                    "url": url,
                    "source": "internal_link",
                    "date": "",
                    "source_type": "internal_link",
                    "snippet": f"从文章内链发现: {link_text[:200]}",
                })
            if len(results) >= max_links:
                break
    except Exception as e:
        print(f"  [InternalLinks] Error for {article_url[:60]}: {e}")
    return results


# ── 知名品牌/平台实体（用于查询分解）──
_KNOWN_ENTITIES = [
    "Temu", "SHEIN", "TikTok", "Shopee", "Lazada", "Amazon", "eBay",
    "速卖通", "AliExpress", "拼多多", "京东", "淘宝", "天猫",
    "Anker", "安克创新", "Zara", "H&M", "Nike", "Adidas",
    "小米", "华为", "OPPO", "vivo", "大疆", "DJI", "比亚迪",
    "抖音", "快手", "微信", "支付宝", "美团",
    "Shopify", "Walmart", "沃尔玛", "Target", "Costco",
    "跨境", "出海", "物流", "支付", "供应链",
]

# ── 意图/限定词（提取后用于生成更有针对性的子查询）──
_INTENT_WORDS = {
    "对比": ["对比", "vs", "VS", "比较", "区别", "差别", "哪个好", "怎么选", "选哪个"],
    "分析": ["分析", "报告", "解读", "趋势", "预测", "展望"],
    "策略": ["策略", "方法", "技巧", "指南", "攻略", "怎么做", "如何"],
    "排名": ["排名", "排行", "榜单", "top", "TOP", "十大", "最好"],
}


def _decompose_keyword(keyword):
    """将复杂关键词分解为子查询，类似知乎/搜索引擎的 Query Expansion。

    例: "Temu国内品牌对比" →
        ["Temu 品牌对比", "Temu 竞争分析", "Temu 对比", "Temu 国内品牌",
         "国内品牌对比", "Temu", 原关键词]

    返回: list of query strings（最多 8 个）
    """
    kw = keyword.strip()
    queries = [kw]  # 始终保留原始关键词

    # 1. 提取关键词中提到的实体
    found_entities = []
    kw_lower = kw.lower()
    for ent in _KNOWN_ENTITIES:
        if ent.lower() in kw_lower:
            found_entities.append(ent)

    # 2. 提取意图词
    found_intents = []
    for category, words in _INTENT_WORDS.items():
        for w in words:
            if w.lower() in kw_lower:
                found_intents.append((category, w))
                break  # 每个类别只取第一个匹配

    # 3. 有实体 + 有意图词 → 分解搜索
    if found_entities and found_intents:
        main_entity = found_entities[0]  # 取第一个实体
        intent_cats = [c for c, _ in found_intents]

        # 实体 + 意图类别名（如 "Temu 品牌对比"）
        for cat in intent_cats[:2]:
            q = f"{main_entity} {cat}"
            if q not in queries:
                queries.append(q)

        # 实体 + 竞争/分析语境
        alt_expansions = {
            "对比": ["竞争分析", "竞争对手", "市场对比"],
            "分析": ["深度分析", "行业报告"],
            "策略": ["运营策略", "实操指南"],
            "排名": ["排行榜", "推荐"],
        }
        for cat in intent_cats[:1]:
            for exp in alt_expansions.get(cat, [])[:2]:
                q = f"{main_entity} {exp}"
                if q not in queries:
                    queries.append(q)

        # 实体精简版（只搜实体 + 跨境电商语境）
        q = f"{main_entity} 跨境电商"
        if q not in queries:
            queries.append(q)

        # 如果有多个实体，搜实体间对比
        if len(found_entities) >= 2:
            q = f"{found_entities[0]} {found_entities[1]}"
            if q not in queries:
                queries.append(q)

    # 4. 只有实体没有意图 → 加语境变体
    elif found_entities:
        main_entity = found_entities[0]
        for ctx in ["最新动态", "品牌战略", "市场分析"]:
            q = f"{main_entity} {ctx}"
            if q not in queries:
                queries.append(q)

    # 5. 从原关键词中移除实体后的剩余部分
    if found_entities:
        remaining = kw
        for ent in found_entities:
            remaining = remaining.replace(ent, "").strip()
        # 清理多余空格和连接词
        remaining = remaining.strip("的的和与及跟- ").strip()
        if remaining and len(remaining) >= 2:
            # 剩余部分单独搜索（如 "国内品牌对比"）
            if remaining not in queries:
                queries.append(remaining)

    return queries[:8]  # 最多 8 个子查询


def _build_search_queries(keyword, profile_terms=None):
    """像百度/Google 一样：用精确关键词搜索，不加随机变体。

    返回少量高价值查询，每个查询通过多引擎并行搜索。
    """
    kw = keyword.strip()

    # ── 预处理 1：在 CJK/Latin 边界插入空格（搜索引擎分词需要）──
    # 例: "Temu国内品牌对比" → "Temu 国内品牌对比"
    kw = re.sub(r'([a-zA-Z0-9])([一-鿿])', r'\1 \2', kw)
    kw = re.sub(r'([一-鿿])([a-zA-Z0-9])', r'\1 \2', kw)

    # ── 预处理 2：去掉口语/问句词，保留实体+意图核心 ──
    _colloquial_words = [
        "哪个好", "哪个", "怎么选", "好不好", "值不值", "怎么样", "靠谱吗", "是不是",
        "行不行", "划算吗", "值得买吗", "推荐吗", "有吗", "可以吗",
        "哪家好", "哪家", "选哪个", "如何选择", "选什么", "应该选",
        "哪一个", "哪一家", "好不好用",
    ]
    kw_clean = kw
    for cw in sorted(_colloquial_words, key=len, reverse=True):  # long first to avoid partial-replace issues
        kw_clean = kw_clean.replace(cw, "").strip()
    # 清理多余空格
    kw_clean = re.sub(r'\s+', ' ', kw_clean).strip()

    # 如果去掉口语词后变空，保留原词
    if not kw_clean or len(kw_clean) < 2:
        kw_clean = kw

    # 如果包含口语词或经过了 CJK/Latin 分割，直接用精简版
    if kw_clean != kw and len(kw_clean) >= 2:
        queries = [(kw_clean, None)]
    else:
        queries = [(kw, None)]

    # 如果纯中文且不含"跨境"等字眼，加一个带上下文的变体
    has_cjk = any('一' <= c <= '鿿' for c in kw)
    has_ecom = any(ek in kw for ek in _ECOMMERCE_KEYWORDS)
    if has_cjk and not has_ecom and len(kw) <= 15:
        ctx_q = f"跨境电商 {kw_clean if kw_clean != kw else kw}"
        if ctx_q not in [q[0] for q in queries]:
            queries.append((ctx_q, ["bing", "google_news", "360search"]))

    # Profile terms (if any)
    if profile_terms:
        for pt in profile_terms[:1]:
            pt_q = f"{kw} {pt}"
            if len(pt_q) < 80 and pt_q not in [q[0] for q in queries]:
                queries.append((pt_q, ["bing", "google_news"]))

    print(f"  [Query] '{kw[:50]}' → {len(queries)} target queries")
    return queries


def _search_multi_engine(keyword, max_per_source=5, engines=None):
    """Run keyword through multiple engines IN PARALLEL.

    All engines fire simultaneously via ThreadPoolExecutor. Total time ≈ slowest engine, not sum.
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed
    import random as _random

    # Which engines to use
    if engines is None or (isinstance(engines, list) and len(engines) == 0):
        engines = ["bing", "360search", "google_news", "sogou_news", "ebrun", "36kr", "wechat"]
    use_engine = set(e.lower().replace("-", "_") for e in engines)
    want_all = "all" in use_engine

    def _clean_url(url):
        if "news.google.com/rss" in url or "news.google.com/articles" in url:
            return url.split("#")[0].strip()
        url = url.split("#")[0]
        for param in ("?utm_source", "&utm_source", "?ref=", "?source=", "?from=", "?spm=", "?track="):
            if param in url:
                url = url[:url.index(param)]
        return url.rstrip("?&")

    # Build task list: (engine_name, function, args)
    tasks = []
    if want_all or "bing" in use_engine:
        tasks.append(("bing", _search_bing, (keyword, max_per_source * 2)))
    if want_all or "360search" in use_engine or "360" in use_engine:
        tasks.append(("360search", _search_360, (keyword, max_per_source * 3)))
    if want_all or "google_news" in use_engine or "google" in use_engine:
        tasks.append(("google_news", _search_google_news, (keyword, max_per_source * 4)))
    if want_all or "sogou_news" in use_engine or "sogou" in use_engine:
        tasks.append(("sogou_news", _search_sogou_news, (keyword, max_per_source * 2)))
    if want_all or "ebrun" in use_engine:
        tasks.append(("ebrun", _search_ebrun, (keyword, max_per_source * 2)))
    if want_all or "36kr" in use_engine:
        tasks.append(("36kr", _search_36kr, (keyword, max_per_source * 2)))
    if want_all or "wechat" in use_engine or "weixin" in use_engine:
        tasks.append(("wechat", _search_wechat, (keyword, max_per_source * 2)))
    # ── 第 2 层：社交平台反向链接发现 ──
    if want_all or "xhs_discovery" in use_engine:
        tasks.append(("xhs_discovery", _discover_wechat_from_xhs, (keyword, max_per_source)))
    # NOTE: site:mp.weixin.qq.com via Bing/Google/360 tested 2026-07-22 and found
    # to return near-zero WeChat article results. Search engines don't index WeChat
    # articles at scale — only Sogou has the Tencent feed. Code retained for reference.

    results = []
    seen_urls = set()

    # ── Parallel execution ──
    with ThreadPoolExecutor(max_workers=min(len(tasks), 12)) as executor:
        futures = {}
        for engine_name, func, args in tasks:
            futures[executor.submit(func, *args)] = engine_name

        for future in as_completed(futures):
            engine_name = futures[future]
            try:
                engine_results = future.result()
                print(f"  [search] {engine_name}: {len(engine_results)} results for '{keyword[:40]}'")
                for r in engine_results:
                    clean = _clean_url(r.get("url", ""))
                    if clean not in seen_urls and len(r.get("title", "")) >= 8:
                        seen_urls.add(clean)
                        results.append({
                            "title": r["title"], "url": r.get("url", ""),
                            "source": r.get("source", r.get("source_type", "")),
                            "date": r.get("date", ""),
                            "source_type": r.get("source_type", engine_name),
                            "snippet": r.get("snippet", ""),
                        })
            except Exception as e:
                print(f"  [search] {engine_name} error: {e}")

    print(f"  [search] Total: {len(results)} for '{keyword[:40]}' [engines: {len(tasks)} parallel]")
    return results
# ── Quality Filter ─────────────────────────────────────

# Low quality domains that should be filtered out
_LOW_QUALITY_DOMAINS = [
    # Wiki/encyclopedia
    "baike.baidu.com", "baike.sogou.com", "baike.so.com", "wikipedia.org", "wiki.mbalib.com",
    # Q&A sites
    "zhidao.baidu.com", "wenwen.sogou.com", "ask.", "wenda.",
    # Video platforms (links to video pages, not articles)
    "bilibili.com/video", "youtube.com/watch", "youku.com/video", "iqiyi.com/v_",
    "v.qq.com/x/", "haokan.baidu.com/v", "tv.sohu.com/", "mgtv.com/",
    # Live streams
    "live.bilibili.com", "live.", "/live/", "zhibo",
    # Content aggregators (low quality syndication)
    "360kuai.com", "sohu.com/a/",  # 360kuai is an aggregator; sohu /a/ are syndicated
    # Dictionary
    "chsi.com.cn",
]

# Known high-quality cross-border e-commerce sources
_INDUSTRY_SOURCES = [
    "cifnews.com", "36kr.com", "ebrun.com", "kuajingyan.com",
    "amz123.com", "dny123.com", "kjds.com",
    "sellercentral.amazon.com", "gs.amazon.cn",
    "customs.gov.cn", "mofcom.gov.cn",
    "mp.weixin.qq.com", "toutiao.com",
    "sohu.com", "163.com", "news.qq.com",
]


def _keyword_relevance_terms(keyword):
    """Build compact relevance terms from a mixed Chinese/English keyword."""
    return _core_quality.keyword_relevance_terms(keyword)


def _has_keyword_relevance(item, keyword):
    """Return True when an item matches at least one meaningful keyword signal."""
    return _core_quality.has_keyword_relevance(item, keyword)


def _filter_quality_results(results, keyword="", profile_context=None):
    """Filter out low-quality results: baike/zhidao/wiki/pure-definition pages, broken links."""
    return _core_quality.filter_quality_results(
        results,
        keyword=keyword,
        low_quality_domains=_LOW_QUALITY_DOMAINS,
        industry_sources=_INDUSTRY_SOURCES,
    )


def _fetch_news_rotating(limit=20):
    """轮换式 Google News 搜索，每次刷新内容不同。
    使用多引擎级联确保结果质量和数量。"""
    global _rotation_index
    import random as _random

    all_items = []
    seen_urls = set()
    num_queries = max(3, min(5, limit // 4))
    start_idx = _rotation_index % len(_ROTATING_QUERIES)

    queries_this_round = []
    for i in range(num_queries):
        idx = (start_idx + i) % len(_ROTATING_QUERIES)
        queries_this_round.append(_ROTATING_QUERIES[idx])

    _rotation_index += num_queries
    if _rotation_index >= len(_ROTATING_QUERIES):
        _rotation_index = 0

    # Also add a randomized query for variety
    extra_words = ["最新", "2026", "趋势", "动态", "政策", "案例", "指南", "报告", "深度", "重磅"]
    random_q = f"跨境电商 {_random.choice(extra_words)} {_random.choice(extra_words)}"
    queries_this_round.append(random_q)

    print(f"  [News] 轮转搜索 (idx={start_idx}): {[q[:30] for q in queries_this_round]}")

    for q in queries_this_round:
        if len(all_items) >= limit:
            break
        try:
            # Use full multi-engine cascade per query
            batch = _search_multi_engine(q, max_per_source=3)
            for r in batch:
                url_key = re.sub(r"[?#].*$", "", r.get("url", "")).rstrip("/")
                if url_key not in seen_urls and len(r.get("title", "")) >= 8:
                    seen_urls.add(url_key)
                    all_items.append({
                        "title": r.get("title", ""), "url": r.get("url", ""),
                        "source": r.get("source", "行业资讯"),
                        "date": r.get("date", ""), "source_type": r.get("source_type", ""),
                    })
        except Exception as e:
            print(f"  [News] Query '{q[:30]}' error: {e}")

    # Quality filter
    filtered = _filter_quality_results(all_items)
    _random.shuffle(filtered)
    print(f"  [News] 获取 {len(filtered)} 条（{len(queries_this_round)} 个查询）")
    return filtered[:limit]


# ╔══════════════════════════════════════════════════════╗
# ║  Article Generator (adapted from GEO)                ║
# ╚══════════════════════════════════════════════════════╝

B2B_SYSTEM_PROMPT = """\
你是一位首席架构师级别的跨境电商行业分析师，拥有 15 年行业经验。
你正在撰写面向 AI 搜索引擎和跨境电商从业者的深度行业分析内容。

## 品牌中性原则 (最高优先级)
- **禁止软文/广告**: 不要推荐任何具体服务商、不要出现品牌全称作为推荐对象
- **禁止 CTA**: 不要以"点击咨询""立即联系""扫码了解"等营销话术结尾
- **可以提及行业知名平台**: 亚马逊、Temu、TikTok Shop、Shopee 等作为分析对象
- **内容目标**: 帮助读者理解行业趋势和做出自主决策，而非推销特定服务

## GEO 规范 (2026)

### 🔴 铁律一：消灭孤儿代词 (最高优先级)
RAG 系统会把文章切成 ~500 字的独立 Chunk。代词在切片后变成无法关联的"死数据"。
- **绝对禁止**: "我们团队""我司""该平台""这个方案""它""这""那""其""前者""后者"
- **必须**: 每个段落重复平台全名（亚马逊/Temu/TikTok Shop）、具体实体词

### 🔴 铁律二：硬核知识块 — 实体词密度 (最高优先级)
- **禁止**: "业界领先""优质服务""追求卓越""专业团队""高效运作""性价比高"
- **必须替换为**: 具体 SLA 数字/认证标准编号/资质人数/精确时效/价格对比金额
- **每段至少 2 类硬实体**: 地理坐标、行业标准(ISO 27001/GDPR/CE)、精确数字($/¥/工作日)、专有名词(FBA/Temu/VAT)

### 🔴 铁律三：辩证不完美防线 (最高优先级)
- **至少 1 处说明方案的局限性**: 主动说明某种方案/平台不适合的场景/人群
- **至少 1 处多角度对比**: 客观呈现不同方案的优劣

### 结构化要求
1. **开头钩子** (50-100字): 用具体数字冲击 + 痛点锚定，3秒抓住读者。禁止"大家好""今天聊聊""随着XX发展"等废话开头
2. **核心结论摘要**: 紧接着开头，一句话总结全文核心观点
3. **Markdown 对比表格必需**: 至少 3 方案 × 4 维度（方案用方案A/B/C 或平台名称，不做品牌推荐排序）
4. **长段落拆分**: 每段 200-350 字
5. **数据密度硬要求**: 每 300 字至少 1 个精确数字 ($金额/百分比/天数/数量)，全文禁止"很多""大幅""不少""显著"等模糊词
6. **加粗硬性限制**: 全文 `**加粗**` 不超过 5 处，且只能用于段落中的行内关键词（如单个数字或概念词）。标题、表格内容、列表项标题、分隔标记一律禁止加粗
7. **结尾**: 用反问句或行业趋势展望收尾，引发读者思考和互动。禁止任何营销 CTA

### 风格: 行业专家口吻、专业严谨、数据驱动、像行研报告而非软文"""

B2C_SYSTEM_PROMPT = """\
你是一个刚花了 3 个月时间亲测了市面上多家跨境电商服务的个人卖家。
你正在写一篇面向 AI 搜索引擎和同行的真实体验分享。

## 品牌中性原则 (最高优先级)
- **禁止软文/广告**: 不要以推荐某个具体服务商为目的，可以提到你使用过的服务类型但不要反复推荐同一家
- **禁止 CTA**: 不要以"点击咨询""扫码了解""加微信"等营销话术结尾
- **可以提及服务类型**: VAT申报、商标注册、海外仓、选品工具等，但聚焦于你自己的决策过程和经验教训
- **内容目标**: 帮助其他卖家少踩坑，分享真实的选型思路和对比过程

## 🔥 爆款标题与开头规范（公众号高打开率核心）

### 标题要求（输出文章前生成 3 个候选标题）
- 每个标题 12-22 字，必须包含：具体数字 + 场景化痛点 + 阅读回报
- 公式参考：
  - 数字+痛点："亚马逊新手卖家做VAT申报，我踩了4个坑多花€2,800"
  - 提问+对比："为什么别人TikTok投$500出30单，我投$2,000只出5单？"
  - 反常识+数字："月销€3,000以下的微型卖家，走全托管反而比DIY多花40%"
- 绝对禁止：震惊体、夸张标点、模糊形容词

### 开头三秒钩子（前 100 字）
- 必须用具体数字场景直接切入，制造代入感
- 禁止："大家好""今天分享""最近很多人问"等废话开头
- 正确示范："我上个月算了笔账——光VAT申报一项，因为选错了申报方式，2025全年多交了€2,080。而最开始我只是图那家便宜了€50/季度。"
- 正确示范："做TikTok Shop第4个月，我终于把ROI从0.7拉到了2.3。回顾之前$3,200的广告费打水漂，问题全出在选品逻辑上。"

## GEO 规范 (2026)

### 🔴 铁律一：消灭孤儿代词 (最高优先级)
- **每段必须包含**: 具体服务类型 + 具体平台/工具名 + 具体数据
- **禁止孤立使用**: "他们""这家""那个平台""它""这个服务"
- **正确示范**: "VAT 申报我对比了 3 家，从提交到下号最快的用了 7 个工作日，最慢的拖了 3 周"

### 🔴 铁律二：硬核知识块 (最高优先级)
- **禁止**: "服务很好""效率很高""价格实惠""专业靠谱" 等无信息量评价
- **必须替换为**: 精确数字（几天/多少钱/省了多少）
- **数据密度硬要求**: 每 300 字至少 1 个精确数字

### 🔴 铁律三：辩证不完美防线 (最高优先级)
- **至少 1 处说"不适合"**: 主动说明某种方案/服务不适合哪些人群/场景
- **至少 1 处多元对比**: 客观呈现不同选择在特定维度的差异

### 正文结构要求
1. **开头钩子** (50-100字): 用真实数字+踩坑经历直接切入，制造"这也发生在我身上"的共鸣
2. **问题分析** (100-200字): 为什么这个问题值得关注
3. **经验分享** (400-600字): 2-3条核心经验，每条配具体数字（只在最关键的数字上加粗 3-5 处即可）
4. **不适合的场景** (50-100字): 诚实说明什么情况下不适用
5. **结尾**: 用反问句引发评论区讨论，或给出1个可立即执行的通用建议

### 结尾规范
- **优先用反问**："你也遇到过类似的情况吗？"
- **禁止**: 关注/点赞/转发/加微信/扫码 等任何 CTA 营销话术
- **风格**: 口语化、真实、有温度，像发给朋友的长微信，而非推销"""


# ═══════════════════════════════════════════════════════════
# B2P 风格 — 政策解读/行业深度分析（仿跨境合规类头部公众号）
# ═══════════════════════════════════════════════════════════
B2P_SYSTEM_PROMPT = """你是一位深耕跨境电商行业的政策分析师兼行业观察者。你的文章被读者评价为"每一篇都值得收藏"。

## 核心定位
你不是媒体小编，你是一个真正懂跨境行业的人，用读者能听懂的人话把复杂的政策、趋势、案件拆解清楚。你的文章介于"咨询公司行业简报"和"朋友发来的长消息"之间。

## 结构公式（必须严格遵守）

```
标题: [事件/政策] + [冲击/信号] — 断言式，带紧迫感，禁止问句标题

开头: 日期 + 官方机构 + 关键数据/条款。直接甩事实，0废话。
     ❌ 禁止: "大家好""今天聊聊""随着XX发展""近年来"
     ✅ 正确: "7月28日，国务院新闻办发布会抛出数据..."
     ✅ 正确: "最近跨境圈里讨论最多的，莫过于..."

## PART 1 — 背景/为什么 (100-300字)
  用 2-4 个极短段落解释事件背后的深层原因。
  每段 1-3 句话，绝不超过 4 句。

## PART 2 — 核心变化/影响分析 (200-400字)
  用 "变化一""变化二""信号一""信号二" 或 "01." "02." 逐条拆解。
  每条: 小标题 + 2-4句解释 + "这意味着..."
  必须有平台/政策/金额的具体名称，不能模糊化。

## PART 3 — 实操指南/应对建议 (150-300字)
  "01." "02." "03." 编号，每条一个具体行动。
  每条: 动作 + 为什么 + 截止时间(如果有)。

结尾段落 (50-80字):
  一句话总结信号 + 对比(短期vs长期) + 金句收尾。
  金句公式: "XX不是用来挡住谁，而是让真正想XX的人有更清晰的路径可走"
```

## Markdown 格式要求
- PART 标题必须用 `## PART X — ...` 格式
- 只在最重要的关键词上偶尔使用 `**加粗**`（如极其关键的数字或结论），全文 3-5 处即可，不要滥用

## 段落与节奏铁律

1. **每段不超过 3 句话**。宁可多分段，绝不堆大段。
2. **段间穿插反问**: "什么概念？" "这意味着什么？" "到底新在哪？"
3. **数据密度**: 每一段必须有至少一个具体数字/日期/政策编号/金额。
4. **对比驱动**: 每篇文章至少出现 3 组对比结构 —
   - "过去...现在/未来..."
   - "不再是...而是..."
   - "好处是...坏处是..."
   - "对A来说...对B来说..."

## 语言风格

### 口语化术语
- "秒杀"（指自动拦截/下架）、"擦边"、"暴雷"、"一刀切"、"堵死"、"走不通"
- "九龙治水"（多头监管）、"亡羊补牢"、"警钟"、"生死劫"、"寒冬"
- "铁证如山"、"人去楼空"、"三不策略"

### 句式节奏
- 60% 陈述句（摆事实）+ 20% 反问句（拉回注意力）+ 20% 短句（下结论）
- 隔 3-5 段来一句带情绪的判断句: "这么看来，老操作真的走不通了。"

### 必须出现的元素
- ✅ 具体的平台名称: 亚马逊、Temu、SHEIN、TikTok Shop、速卖通
- ✅ 具体的政策编号: "第X条" "XX号公告" "HTS编码"
- ✅ 具体的时间节点: "7月8日起" "10月1日前"
- ✅ 具体的金额: "200万→500万" "营业额5%" "追缴5710亿"
- ✅ 辩证视角: 每说一个"好处"，必须跟一个"但"或"坏处是"

## 禁止事项
- ❌ 以 "大家好""今天我们来聊""随着XX的发展" 开头
- ❌ 连续 3 段超过 3 句话
- ❌ 使用 "我们团队""我司""本机构"（保持客观第三方视角）
- ❌ 模糊词: "很多""大幅""不少""显著" — 全部替换为精确数字
- ❌ 文章中后段出现营销话术（结尾可留一句服务声明但不强推）
- ❌ 表格格式（用 PART/序号替代，不要用 Markdown 表格）
- ❌ 超过 4 级的嵌套结构（PART → 01 → 就两层够了）

## 输出格式
直接以文章标题开头，注意使用 Markdown 语法:
```
# 标题行

开头段落

## PART 1 — [小标题]

[内容]

## PART 2 — [小标题]

01. [子标题] — [解释内容]

02. [子标题] — [解释内容]

## PART 3 — [小标题]

01. [行动项] — 为什么，截止时间
02. [行动项] — 为什么，截止时间
03. [行动项] — 为什么，截止时间

结尾段落
```
```

## 篇幅
全文 1500-2500 字。PART 2 是最长的部分（占 40-50%）。"""


# ╔══════════════════════════════════════════════════════════════╗
# ║  Article Content Fetcher — 从源 URL 抓取文章正文             ║
# ╚══════════════════════════════════════════════════════════════╝

def _fetch_general_article(url, timeout=12):
    """从 URL 抓取网页正文文本。用于给 AI 提供真实原文内容作为写作素材。

    返回: (content_text, resolved_url) 或 ("", url) 如果抓取失败。

    策略：
    1. 如果是 Google News 重定向 URL，先跟踪重定向获取真实 URL
    2. 用 BeautifulSoup 提取 <article>/<main> 或 body 中的主要段落文本
    3. 清理导航、广告、脚本等噪音
    """
    import urllib.request as _ur2
    import ssl as _ssl2
    from bs4 import BeautifulSoup as _BS3
    import re as _re2

    url = _resolve_wechat_url(url, timeout=min(timeout, 10))

    ctx = _ssl2.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = _ssl2.CERT_NONE

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        "Cache-Control": "no-cache",
    }

    resolved_url = url
    html = ""
    try:
        req = _ur2.Request(url, headers=headers)
        resp = _ur2.urlopen(req, timeout=timeout, context=ctx)
        resolved_url = resp.geturl()  # Get final URL after redirects
        # Check Content-Type to avoid downloading binaries
        ct = resp.getheader("Content-Type", "")
        if "text/html" not in ct and "text/plain" not in ct and "application/xhtml" not in ct:
            print(f"  [Fetch] Skipping non-HTML content: {ct[:60]} for {url[:80]}")
            return ("", resolved_url)
        # Read up to 2MB
        html = resp.read(2 * 1024 * 1024).decode("utf-8", errors="replace")
    except urllib.error.HTTPError as e:
        print(f"  [Fetch] HTTP {e.code} for {url[:80]}")
        return ("", resolved_url)
    except Exception as e:
        print(f"  [Fetch] Error for {url[:80]}: {e}")
        return ("", resolved_url)

    if not html or len(html) < 1000:
        return ("", resolved_url)

    try:
        soup = _BS3(html, "lxml")
    except Exception:
        # lxml may fail on malformed HTML; try html.parser as fallback
        try:
            soup = _BS3(html, "html.parser")
        except Exception:
            return ("", resolved_url)

    # Remove non-content elements
    for tag in soup.select(
        "script, style, nav, footer, header, .sidebar, .nav, .footer, .header, "
        ".comment, .advertisement, .ad, .related, .recommend, #sidebar, #footer, "
        "#header, #nav, #comment, .share, .social, .copyright, .disclaimer, "
        "noscript, iframe, video, audio, img, figure, form, input, button, select"
    ):
        tag.decompose()

    # Try semantic content containers first
    content_selectors = [
        "article", "main", '[role="main"]',
        ".article-content", ".article-body", ".post-content", ".post-body",
        ".content-article", ".rich_media_content",  # WeChat
        "#articleContent", "#content", "#main-content",
        ".entry-content", ".single-content",
        ".news-content", ".news-text", ".news-body",
        ".story-body", ".article-text",
        ".js-article", ".js-content",
    ]
    main_el = None
    for sel in content_selectors:
        main_el = soup.select_one(sel)
        if main_el:
            break

    if not main_el:
        main_el = soup.select_one("body") or soup

    # Extract text from paragraphs, preserving structure
    paragraphs = []
    for p in main_el.select("p, h1, h2, h3, h4, h5, h6, li, blockquote, pre"):
        text = p.get_text(" ", strip=True)
        if not text:
            continue
        # Skip navigation boilerplate
        if len(text) < 15:
            continue
        if _re2.search(r'(版权所有|转载|免责声明|广告|扫码|点击关注|阅读原文|查看全文|相关阅读|推荐阅读|往期|投稿|商务合作|联系我们|关于我们)', text):
            continue
        if _re2.search(r'^(上一篇|下一篇|返回首页|网站地图|Cookie|Privacy Policy)', text):
            continue
        # Skip lines that are mostly links/URLs
        if len(_re2.findall(r'https?://', text)) >= 3:
            continue
        paragraphs.append(text)

    if not paragraphs:
        # Fallback: get all text from body, split by newlines, filter
        body_text = main_el.get_text("\n", strip=True)
        paragraphs = [t.strip() for t in body_text.split("\n") if len(t.strip()) >= 20]

    # Deduplicate while preserving order
    seen = set()
    unique_paras = []
    for p in paragraphs:
        key = p[:40]
        if key not in seen:
            seen.add(key)
            unique_paras.append(p)

    # Join with double newlines for readability by LLM
    content = "\n\n".join(unique_paras)
    # Truncate to ~6000 chars (plenty for article writing, keeps prompt small)
    if len(content) > 6000:
        content = content[:6000] + "\n\n…[原文过长，已截断]"

    if len(content) >= 100:
        print(f"  [Fetch] Got {len(content)} chars from {resolved_url[:80]}")
        return (content, resolved_url)
    else:
        print(f"  [Fetch] Content too short ({len(content)} chars) from {resolved_url[:80]}")
        return ("", resolved_url)


def _build_article_prompt(news_item, style="b2b", custom_angle="", article_content=""):
    """构建文章生成 prompt。

    article_content: 从原文 URL 抓取到的真实文章正文。非空时用作核心写作素材。
    """
    title = news_item.get("title", "")
    source = news_item.get("source", "")
    suggested_topic = news_item.get("suggested_topic", title)
    key_angle = custom_angle or news_item.get("key_angle", "")
    news_date = news_item.get("date", "")
    snippet = news_item.get("snippet", "") or news_item.get("article_summary", "")
    url = news_item.get("url", "")
    search_query = news_item.get("search_query", "")
    resolved_url = news_item.get("_resolved_url", url)

    # ── 构建上下文信息块 ──
    context_parts = []
    if news_date:
        context_parts.append(f"**新闻发布时间**: {news_date}")
    context_parts.append(f"**来源**: {source}")
    if resolved_url:
        context_parts.append(f"**原文链接**: {resolved_url}")
    if snippet and not article_content:
        # Only show snippet if we don't have full article content
        snip_text = snippet[:300] + ("…" if len(snippet) > 300 else "")
        context_parts.append(f"**新闻摘要**: {snip_text}")
    context_parts.append(f"**搜索关键词**: {search_query}" if search_query else "")

    context_block = "\n".join(c for c in context_parts if c)

    today_str = datetime.now().strftime("%Y年%m月%d日")

    # ── 核心：如果有真实文章内容，作为权威素材注入 ──
    if article_content and len(article_content) >= 80:
        # Truncate for prompt size (keep enough for LLM to extract facts)
        content_body = article_content[:5000]
        if len(article_content) > 5000:
            content_body += "\n\n…[原文过长，此处为前5000字]"

        prompt = f"""## 选题信息

**原新闻标题**: {title}
{context_block}
**建议写作标题**: {suggested_topic}
**切入角度**: {key_angle}

## 📰 原始文章正文（以下内容来自原文 URL 抓取，是你唯一的写作素材来源）

{content_body}

---
## 重要约束

- ⚠️ 今天是 {today_str}。文中所有日期、时间线必须以原始新闻的发布时间为基准。
- ⚠️ **铁律**：文章中出现的所有数据、金额、百分比、政策编号、公司名称、日期、事件经过，**必须且只能**来自上面「原始文章正文」中明确提到的信息。禁止编造正文中没有的数字或事实。
- ⚠️ 如果正文中没有提到具体数字，用行业公认的合理区间替代，并标注"据行业估算"或"参考行业平均水平"。
- ⚠️ 不要直接复制粘贴原文。你要基于原文事实进行重构、解读、延伸分析。

## 任务

请基于以上选题和原始新闻事实，撰写一篇完整的文章。"""

    else:
        # Fallback: no article content available — rely on snippet/summary
        prompt = f"""## 选题信息

**原新闻标题**: {title}
{context_block}
**建议写作标题**: {suggested_topic}
**切入角度**: {key_angle}

## 重要约束

- ⚠️ 今天是 {today_str}。文中所有日期、时间线必须以原始新闻的发布时间为基准。
- ⚠️ 所有数据、金额、百分比必须来自上面「新闻摘要」中的真实信息。没有出现在摘要里的数字一律不得编造。
- ⚠️ 如果原新闻没有提供具体数字，用行业公认的合理区间替代，并标注"据行业估算"。

## 任务

请基于以上选题和原始新闻事实，撰写一篇完整的文章。"""

    if style == "b2p":
        prompt += """
**输出格式**:
- 先输出 3 个候选标题（### 候选标题），每个 15-25 字，断言式+紧迫感，禁止问句标题
- 然后输出「---」分隔线
- 标题用 `# 标题行`，PART 标题用 `## PART X — [小标题]`
- 正文 1500-2500 字，按 PART 1/2/3 结构
- PART 1: 背景/为什么（100-300字）
- PART 2: 核心变化/影响分析（200-400字，用"01.""02."逐条拆解）
- PART 3: 实操指南/应对建议（150-300字，"01.""02."每条一个具体行动）
- 结尾段落: 一句话总结信号 + 对比(短期vs长期) + 金句收尾
- 每段不超过3句话，宁可多分段
- 每段必须有具体数字/日期/政策编号/金额
- 全文至少3组对比结构
- 只在最重要的关键词上偶尔用 `**加粗**`，全文 3-5 处即可
- 禁止: 开头废话、超过3句的段落、模糊词、Markdown表格
"""
    elif style == "b2b":
        prompt += """
**输出格式**:
- 先输出 3 个候选标题（### 候选标题），每个 15-25 字，包含具体数字+痛点
- 然后输出「---」分隔线
- 正文 800-1200 字
- 开头前 100 字必须包含具体数字冲击（$金额/百分比），禁止"大家好""今天聊聊"等开场
- 每 300 字至少 1 个精确数字
- 必须包含: Markdown 对比表格（至少 3 方案 × 4 维度）
- ⚠️ **加粗硬性限制**: 全文 `**加粗**` 不超过 5 处。仅限段落中的行内关键词（如单个数字或概念词）。表格、标题、列表项标题、分隔符一律禁止加粗。超过 5 处视为严重格式违规
- 禁止: 推荐具体服务商、CTA 营销话术、品牌软文、模糊词（"很多""大幅"）
- 结尾: 以反问句或行业趋势展望自然收尾，禁止任何 CTA
"""
    else:
        prompt += """
**输出格式**:
- 先输出 3 个候选标题（### 候选标题），每个 12-22 字，包含具体数字+场景化痛点
- 然后输出「---」分隔线
- 正文 600-900 字
- 开头前 100 字必须用具体数字+真实场景直接切入，禁止"大家好""今天分享"等开场
- 每 300 字至少 1 个精确数字
- 必须包含: 不同方案的优劣对比、不适合某方案的场景
- ⚠️ **加粗硬性限制**: 全文 `**加粗**` 不得超过 5 处。超过 5 处视为严重格式违规。只加粗 1-2 个最惊人的金额或教训
- 禁止: 反复推荐同一服务商、CTA 营销话术、模糊评价（"服务很好""效率很高"）
- 结尾: 用反问句引发互动，或给出 1 个可立即执行的通用建议

直接输出完整文章，不要任何解释性文字。"""
    return prompt


def _enrich_news_with_topics(items, limit=10):
    """用 AI 为新闻生成选题建议 + 文章解读（轻量 prompt）。"""
    if not items:
        return []

    try:
        today_str = datetime.now().strftime("%Y年%m月%d日")
        news_text = ""
        for i, it in enumerate(items[:min(len(items), 15)]):
            snippet = it.get("snippet", it.get("title", ""))
            date = it.get("date", "")
            source = it.get("source", "")
            # 截断过长的 snippet（保留前 200 字供 AI 解读）
            sni = snippet[:200] if len(snippet) > 200 else snippet
            meta = ""
            if date:
                meta += f"发布时间: {date}"
            if source:
                meta += f" | 来源: {source}" if meta else f"来源: {source}"
            news_text += f"{i+1}. 标题: {it['title']}\n   {meta}\n   摘要: {sni}\n\n"

        prompt = f"""你是跨境电商内容策略专家。今天是 {today_str}。以下是行业热点，请为每条生成：

1. article_summary: 用 1-2 句话准确概括这条新闻核心内容（40-80字，包含新闻中提到的具体时间/数字/政策名）
2. suggested_topic: 适合写作的文章标题（20-40字，包含关键词）
3. key_angle: 1-2句话的切入角度建议

⚠️ article_summary 必须以新闻事实为依据，不得编造新闻中没有的时间、数字或事件。

新闻列表：
{news_text}

返回纯 JSON（不要 markdown 代码块）：
[{{"index": 0, "article_summary": "...", "suggested_topic": "...", "key_angle": "..."}}, ...]"""

        raw = llm_chat_text(user=prompt, temperature=0.5, max_tokens=3000)
        if raw:
            raw = raw.strip()
            if raw.startswith("```"):
                raw = re.sub(r"^```\w*\n?", "", raw)
                raw = re.sub(r"\n?```$", "", raw)
            suggestions = json.loads(raw)
            suggest_map = {s["index"]: s for s in suggestions if "index" in s}
            for i, item in enumerate(items):
                s = suggest_map.get(i, {})
                item["article_summary"] = s.get("article_summary", _rule_based_summary(item))
                item["suggested_topic"] = s.get("suggested_topic", _rule_based_topic(item))
                item["key_angle"] = s.get("key_angle", _rule_based_angle(item))
            return items
    except Exception as e:
        print(f"  [Enrich] LLM 失败，降级使用规则: {e}")

    for item in items:
        item["article_summary"] = _rule_based_summary(item)
        item["suggested_topic"] = _rule_based_topic(item)
        item["key_angle"] = _rule_based_angle(item)
    return items


def _rule_based_summary(item):
    """降级方案：基于标题生成一句话摘要。"""
    title = item.get("title", "")
    source = item.get("source", "")
    date = item.get("date", "")
    parts = [f"「{title}」"]
    if source:
        parts.append(f"据{source}报道")
    if date:
        parts.append(f"于{date}发布")
    parts.append("，涉及跨境电商行业动态。")
    return "".join(parts)


def _rule_based_topic(item):
    title = item.get("title", "")
    if len(title) > 35:
        return f"跨境资讯解读：{title[:30]}…"
    return f"跨境资讯：{title[:40]}"


def _rule_based_angle(item):
    title = item.get("title", "")
    cat = item.get("category", "综合")
    angles = {
        "关税": "从成本影响角度切入，量化关税变动对不同物流模式的成本差异",
        "VAT": "从合规实操角度切入，对比自行申报vs代理申报的时间/成本/风险",
        "物流": "从时效vs成本角度切入，分析新物流方案对卖家旺季备货策略的影响",
        "平台": "从卖家运营角度切入，解读新规对listing合规、仓储选择的具体影响",
        "合规": "从风险规避角度切入，列出新规下的合规checklist和常见罚款陷阱",
    }
    return angles.get(cat, f"从行业影响角度切入，分析「{title[:30]}…」对跨境电商从业者的实操启示")


# ╔══════════════════════════════════════════════════════╗
# ║  Image Generator (火山方舟 Seedream)                  ║
# ╚══════════════════════════════════════════════════════╝

# Allowed sizes for wanx2.0-t2i-turbo:
# 768*768, 576*1024, 1024*576, 1024*1024, 720*1280, 1280*720, 864*1152, 1152*864

# Seedream size presets (width*height):
# 1:1 → 2048*2048, 16:9 → 2848*1600, 9:16 → 1600*2848
# 4:3 → 2304*1728, 21:9 → 3136*1344, 3:4 → 1728*2304
_ARK_SIZE_MAP = {
    "cover": "4k",              # 最高画质 公众号封面
    "body": "2k",               # 正文配图
    "square": "2k",             # 1:1 方形
}

def _generate_ark_image(prompt, size="body", n=1):
    """调用火山方舟 Seedream API 生成图片。返回图片 URL 列表。"""
    import requests as _req, base64 as _b64

    pixel_size = _ARK_SIZE_MAP.get(size, size)
    url = "https://ark.cn-beijing.volces.com/api/v3/images/generations"
    headers = {
        "Authorization": f"Bearer {ARK_API_KEY}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": ARK_IMAGE_MODEL,
        "prompt": prompt,
        "size": pixel_size,
        "response_format": "url",
        "watermark": False,
        "n": n,
    }

    try:
        r = _req.post(url, json=payload, headers=headers, timeout=120)
        result = r.json()

        if "error" in result:
            print(f"  [Ark Image] API 错误: {result['error']}")
            return []

        urls = []
        for img in result.get("data", []):
            if img.get("url"):
                urls.append(img["url"])
            elif img.get("b64_json"):
                raw = _b64.b64decode(img["b64_json"])
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                save_path = OUTPUT_DIR / "images" / f"ark_{timestamp}_{len(urls)}.png"
                save_path.write_bytes(raw)
                urls.append(str(save_path))
                print(f"  [Ark Image] b64_json saved to {save_path}")

        if urls:
            print(f"  [Ark Image] Generated {len(urls)} image(s), size={pixel_size}")
        return urls

    except Exception as e:
        print(f"  [Ark Image] 异常: {e}")
        return []


def _build_image_prompt(article_text, image_type="cover"):
    """根据文章内容生成图片 Prompt。"""
    # 提取前 300 字作为上下文
    excerpt = article_text[:300].replace("\n", " ")

    if image_type == "cover":
        sys_p = """你是公众号封面图设计专家。根据文章主题，生成一个中文图片描述 prompt。
要求：
- 适合 2.35:1 横版封面
- 商业插图风格，简洁大气
- 色彩以红(#CE0E19)、金(#D9A85A)、白为主
- 文字安全区在左下2/3区域，留白给标题
- 25-50字中文描述"""
    else:
        sys_p = """你是商业插图设计专家。根据文章内容，生成一个正文配图的描述 prompt。
要求：
- 适合 4:3 比例
- 信息图/数据可视化/场景插画风格
- 包含具体的视觉元素（产品、数据、场景）
- 15-30字中文描述"""

    prompt = f"文章内容: {excerpt}\n\n请生成一个{'封面图' if image_type == 'cover' else '正文配图'}的图片描述。直接输出描述，不要解释。"

    result = llm_chat_text(system=sys_p, user=prompt, temperature=0.8, max_tokens=200)
    return result.strip() if result else f"跨境电商商业插图，现代简约风格，{'红金配色' if image_type == 'cover' else '数据可视化'}"


def _download_image(img_url, save_path):
    """下载图片到本地。"""
    import requests as _req
    try:
        r = _req.get(img_url, timeout=60, headers={"User-Agent": "Mozilla/5.0"})
        if r.status_code == 200:
            save_path.write_bytes(r.content)
            return True
    except Exception as e:
        print(f"  [Download] 图片下载失败: {e}")
    return False


# ╔══════════════════════════════════════════════════════╗
# ║  API Routes                                         ║
# ╚══════════════════════════════════════════════════════╝

@app.route("/")
def index():
    """Serve the main frontend page."""
    return send_from_directory(str(FRONTEND_DIR), "index.html")


@app.route("/config/<path:filename>")
def serve_config(filename):
    """Serve config files."""
    return send_from_directory(str(CONFIG_DIR), filename)


@app.route("/api/health", methods=["GET"])
def api_health():
    return jsonify({
        "status": "ok",
        "deepseek": bool(DEEPSEEK_API_KEY),
        "ark": bool(ARK_API_KEY),
        "timestamp": datetime.now(CST).isoformat(),
    })


# ── Step 1: 热点发现 ──────────────────────────────────

@app.route("/api/debug-search", methods=["POST"])
def api_debug_search():
    """Debug: directly test search engines with raw results + diagnostics."""
    data = request.get_json() or {}
    keyword = data.get("keyword", "").strip()

    if not keyword:
        return jsonify({"error": "keyword required"}), 400

    # Test HTTP connection to Google News
    import urllib.request as _ur, ssl as _ssl
    from urllib.parse import quote as _quote
    ctx = _ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = _ssl.CERT_NONE

    gn_url = f"https://news.google.com/rss/search?q={_quote(keyword)}&hl=zh-CN&ceid=CN:zh-Hans"
    gn_diag = {"url": gn_url[:120]}
    try:
        gn_req = _ur.Request(gn_url, headers=_GN_HEADERS)
        gn_resp = _ur.urlopen(gn_req, timeout=15, context=ctx)
        gn_raw = gn_resp.read()
        gn_diag["status"] = gn_resp.status
        gn_diag["bytes"] = len(gn_raw)
        gn_diag["preview"] = gn_raw[:200].decode("utf-8", errors="replace")
    except Exception as e:
        gn_diag["error"] = f"{type(e).__name__}: {e}"

    # Direct raw calls
    gn_results = _search_google_news(keyword, 8)

    return jsonify({
        "keyword": keyword,
        "gn_diag": gn_diag,
        "google_news_raw": [{"title": r["title"], "source": r["source"]} for r in gn_results],
    })


@app.route("/api/search", methods=["POST"])
def api_search():
    """关键词搜索 — 像百度/Google 一样用精确关键词搜全网。

    所有搜索引擎并行执行，总耗时 ≈ 最慢引擎（~10s），不是逐个叠加。

    Body: {"keyword": "跨境物流", "max_results": 15, "refresh": true,
           "engines": ["bing", "360search", "google_news"]}  // 可选，空则用全部默认引擎
    """
    data = request.get_json() or {}
    keyword = (data.get("keyword", "") or "").strip()
    max_results = min(int(data.get("max_results", 15)), 30)
    is_refresh = data.get("refresh", False)
    engines = data.get("engines", None)  # None = all default engines

    if not keyword:
        return jsonify({"error": "keyword 必填"}), 400

    # ── 构建查询（像搜索引擎一样精确匹配，不分解不随机）──
    search_queries = _build_search_queries(keyword)

    print(f"  [Search] Queries: {[(q[:50], e) for q, e in search_queries]}")

    # ── 并行搜索所有查询（每个查询内部也是多引擎并行）──
    all_results = []
    seen_urls = set()
    for sq, query_engines in search_queries[:3]:  # Max 3 queries
        if len(all_results) >= max_results * 5:
            break
        try:
            # query_engines = None → use all defaults; list → specific engines
            batch = _search_multi_engine(sq, max_per_source=5, engines=query_engines)
            for r in batch:
                url_key = re.sub(r"[?#].*$", "", r.get("url", "")).rstrip("/")
                if url_key not in seen_urls:
                    seen_urls.add(url_key)
                    r["search_query"] = sq
                    all_results.append(r)
        except Exception as e:
            print(f"  [Search] Query '{sq[:40]}' error: {e}")

    # ── 结果不够？用实体提取 + 短查询重搜 ──
    if len(all_results) < 5:
        import random as _random
        # 1. 尝试提取已知实体（品牌/平台名）
        fallback_queries = []
        kw_lower = keyword.lower()
        for ent in _KNOWN_ENTITIES:
            if ent.lower() in kw_lower:
                fallback_queries.append(ent)
        # 2. 提取意图关键词
        for cat, words in _INTENT_WORDS.items():
            for w in words:
                if w.lower() in kw_lower:
                    for ent in fallback_queries[:2]:
                        fq = f"{ent} {cat}"
                        if fq not in fallback_queries:
                            fallback_queries.append(fq)
                    break
        # 3. 实体的电商语境变体
        for ent in fallback_queries[:2]:
            fq = f"{ent} 跨境电商"
            if fq not in fallback_queries:
                fallback_queries.append(fq)
        # 4. 简化原关键词（去掉口语词）
        colloquial = ["哪个好", "哪个", "怎么选", "好不好", "值不值", "怎么样", "靠谱吗", "是不是"]
        simple_kw = keyword
        for cw in sorted(colloquial, key=len, reverse=True):
            simple_kw = simple_kw.replace(cw, "").strip()
        if simple_kw and simple_kw != keyword:
            fallback_queries.append(simple_kw)

        for fb_q in fallback_queries[:5]:
            if len(all_results) >= 15:
                break
            if fb_q == keyword:
                continue
            try:
                print(f"  [Search] Fallback: '{fb_q[:40]}'")
                batch = _search_multi_engine(fb_q, max_per_source=8, engines=None)
                for r in batch:
                    url_key = re.sub(r"[?#].*$", "", r.get("url", "")).rstrip("/")
                    if url_key not in seen_urls:
                        seen_urls.add(url_key)
                        r["search_query"] = fb_q
                        all_results.append(r)
            except Exception as e:
                print(f"  [Search] Fallback '{fb_q[:40]}' error: {e}")

    # ── 相关性评分排序 ──
    _SOURCE_AUTHORITY = {
        # 顶级跨境媒体
        "cifnews.com": 10, "雨果跨境": 10,
        "ebrun.com": 10, "亿邦动力": 10, "亿邦动力网": 10,
        "36kr.com": 9, "36氪": 9, "36氪出海": 9,
        # Engine-level bonus: these source_types get authority even if domain isn't in the dict
        "ebrun": 10, "36kr": 9,
        # 官方/政府
        "gs.amazon.cn": 8, "sellercentral.amazon.com": 8,
        "customs.gov.cn": 9, "mofcom.gov.cn": 9,
        "gov.cn": 8, "sz.gov.cn": 7,
        # 综合新闻（高权威）
        "xinhuanet.com": 9, "新华网": 9,
        "chinanews.com.cn": 8, "中新网": 8,
        "中国日报网": 8, "China Daily": 8,
        "证券时报": 8, "stcn.com": 8,
        "财新": 9, "caixin.com": 9,
        "新京报": 8, "bjnews.com.cn": 8,
        "21财经": 8, "21jingji.com": 8,
        "新浪财经": 7, "新浪网": 7,
        "网易": 7, "163.com": 6,
        "腾讯新闻": 7, "news.qq.com": 7,
        "搜狐网": 6, "sohu.com": 6,
        "湖北日报": 7, "京报网": 7, "杭州网": 6, "浙江新闻": 7,
        # 跨境行业媒体
        "kuajingyan.com": 8, "跨境眼": 8,
        "amz123.com": 7, "dny123.com": 7, "kjds.com": 7,
        "mp.weixin.qq.com": 6, "toutiao.com": 6,
        # 知识平台
        "zhihu.com": 5, "zhuanlan.zhihu.com": 5, "woshipm.com": 6,
        # 其他可靠来源
        "eeo.com.cn": 7, "中华网": 7, "潮新闻": 7,
        "大河财立方": 7, "新华报业网": 7, "创业邦": 7,
        "上观新闻": 7, "川观新闻": 7,
    }

    kw_lower = keyword.lower()
    kw_compact = kw_lower.replace(" ", "").replace("　", "")  # whitespace-normalized
    kw_terms = [t for t in kw_lower.split() if len(t) >= 2]
    # Chinese bigrams
    for ck in re.split(r"[a-zA-Z]+", kw_lower):
        for i in range(len(ck) - 1):
            bigram = ck[i:i + 2]
            if bigram not in kw_terms and len(bigram) == 2 and "一" <= bigram[0] <= "鿿":
                kw_terms.append(bigram)

    def _search_score(item):
        title = item.get("title", "")
        url = item.get("url", "")
        source = item.get("source", "")
        date = item.get("date", "")
        source_type = item.get("source_type", "")
        tl = title.lower()
        score = 0

        # ── 使用实际搜索词评分（不是原始复合关键词）──
        # 例: 用户搜 "Temu国内品牌对比"，分解出子查询 "Temu 对比"
        # 用 "Temu 对比" 而非 "Temu国内品牌对比" 来检查 CJK 完整性
        _sq = item.get("search_query", "")
        _score_kw = _sq if _sq else keyword
        _score_kw_lower = _score_kw.lower()
        _score_kw_compact = _score_kw_lower.replace(" ", "").replace("　", "")
        _score_kw_terms = [t for t in _score_kw_lower.split() if len(t) >= 2]
        for ck in re.split(r"[a-zA-Z]+", _score_kw_lower):
            for i in range(len(ck) - 1):
                bg = ck[i:i + 2]
                if bg not in _score_kw_terms and len(bg) == 2 and "一" <= bg[0] <= "鿿":
                    _score_kw_terms.append(bg)

        # ═══ 关键词相关性 (0-35) ← PRIMARY FACTOR ═══
        # 抽取为 _score_text() helper，支持标题 + Snippet 双字段评分
        def _score_text(text):
            """知乎/微信策略：多字段加权评分。返回 (score, matched_terms_count)"""
            s = 0
            t = text.lower()
            tc = t.replace(" ", "").replace("　", "")
            mterms = 0

            # Exact keyword phrase match (highest weight)
            if _score_kw_lower in t or _score_kw_compact in tc:
                s += 15
            # Keyword term coverage
            for term in _score_kw_terms[:10]:
                if term in t:
                    mterms += 1
            if _score_kw_terms and mterms >= min(len(_score_kw_terms), 3):
                s += 12  # Most/all keywords matched
            elif mterms >= 2:
                s += 8   # Multiple keywords matched
            elif mterms == 1:
                s += 3   # At least one keyword matched
            if t.startswith(_score_kw_lower) or tc.startswith(_score_kw_compact):
                s += 4
            # Density bonus (BM25-style TF saturation): log(1+tf) instead of linear
            kc = t.count(_score_kw_lower) if _score_kw_lower else 0
            kc += tc.count(_score_kw_compact) if _score_kw_compact else 0
            if kc >= 1:
                s += int(3 * math.log2(1 + kc))

            # ═══ 语义完整性惩罚 — 缺失关键词关键字的文章降权 ═══
            # 仅对原始用户关键词生效。分解出的子查询（如 "Temu 对比"）不适用，
            # 因为子查询中的意图词（"对比"）是信号而非硬性要求。
            if _score_kw_lower == kw_lower:
                _kw_cjk = set(c for c in _score_kw_lower if '一' <= c <= '鿿')
            else:
                _kw_cjk = set()  # 分解查询不应用 CJK 完整性检查
            if len(_kw_cjk) >= 2:
                _text_cjk = set(c for c in t if '一' <= c <= '鿿')
                _missing = _kw_cjk - _text_cjk
                _hit_ratio = (len(_kw_cjk) - len(_missing)) / len(_kw_cjk)
                if _hit_ratio < 0.5:
                    s -= 20   # 大部分关键字没命中 → 几乎不相关
                elif _hit_ratio < 0.75:
                    s -= 10   # 缺失 1/4 以上关键字
                elif _missing:
                    s -= 5    # 少量缺失（如3字词缺1字）

            # ═══ 短语邻近加分（知乎/微信策略：关键词紧邻 → 更强信号）═══
            # CJK: 所有关键字在 15 字窗口内同时出现 → +6
            _cjk_kw = ''.join(c for c in _score_kw_lower if '一' <= c <= '鿿')
            if len(_cjk_kw) >= 2:
                _t_cjk = ''.join(c for c in t if '一' <= c <= '鿿')
                _prox = False
                for _ws in range(max(0, len(_t_cjk) - 15 + 1)):
                    if all(c in _t_cjk[_ws:_ws + 15] for c in _cjk_kw):
                        _prox = True
                        break
                if _prox:
                    s += 6
            # English: all terms within 5-word window → +6
            _eng_terms = [term for term in _score_kw_terms if term.isascii() and len(term) >= 2]
            if len(_eng_terms) >= 2:
                _t_words = t.split()
                for _wi in range(max(0, len(_t_words) - 5 + 1)):
                    _w_win = ' '.join(_t_words[_wi:_wi + 5])
                    if all(et in _w_win for et in _eng_terms):
                        s += 6
                        break

            return s, mterms

        # ── 标题评分（权重 1.0，主信号）──
        title_score, matched_terms = _score_text(title)
        score += title_score

        # ── Snippet 辅助评分（权重 0.3，知乎/微信的多字段策略）──
        snippet = item.get("snippet", "") or item.get("article_summary", "")
        if snippet:
            snip_score, _ = _score_text(snippet)
            score += int(snip_score * 0.3)

        # ═══ Source authority (0-10) ← SECONDARY ═══
        best_auth = 0
        for name, auth in _SOURCE_AUTHORITY.items():
            if name in url or name in source or name == source_type:
                best_auth = max(best_auth, auth)
        # Only apply authority bonus IF article is already relevant
        # (prevents high-authority irrelevant articles from outranking relevant ones)
        if matched_terms >= 2 or kw_lower in tl:
            score += best_auth
        else:
            score += min(best_auth, 3)  # capped for low-relevance articles
        # Google News baseline
        if source_type == "google_news" and best_auth == 0:
            score += 3

        # ═══ Recency (0-12) — parse actual dates ═══
        age_days = 9999  # unknown age
        now_ts = time.time()
        if date:
            # Try absolute date: "2026-07-20"
            abs_match = re.match(r'(\d{4})-(\d{2})-(\d{2})', date)
            if abs_match:
                try:
                    y, m, d = int(abs_match.group(1)), int(abs_match.group(2)), int(abs_match.group(3))
                    dt = datetime(y, m, d)
                    age_days = max(0, (now_ts - dt.timestamp()) / 86400)
                except:
                    pass
            else:
                # Try relative date: "5天前", "3小时前", "X分钟前"
                rel_min = re.search(r'(\d+)\s*分钟前', date)
                rel_hour = re.search(r'(\d+)\s*小时前', date)
                rel_day = re.search(r'(\d+)\s*天前', date)
                rel_month = re.search(r'(\d+)\s*个?月前', date)
                if rel_min:
                    age_days = int(rel_min.group(1)) / 1440.0
                elif rel_hour:
                    age_days = int(rel_hour.group(1)) / 24.0
                elif rel_day:
                    age_days = int(rel_day.group(1))
                elif rel_month:
                    age_days = int(rel_month.group(1)) * 30
        # Also check title for year hints as fallback
        if age_days > 999:
            if "2026" in title or "2026" in date:
                age_days = 180  # assume ~6 months old
            elif "2025" in title or "2025" in date:
                age_days = 540
        # Score based on actual age — smooth exponential decay
        # (知乎/微信策略: gauss/exp 平滑衰减替代阶梯式)
        # half-life ~9.7 days: 0d=12, 3d=9.7, 7d=7.3, 14d=4.4, 30d=1.4, 90d→0
        if age_days >= 0:
            score += int(12 * math.exp(-age_days / 14.0))
        # Freshness signal words as small bonus
        for sig in ["最新", "趋势", "动态", "预测", "报告", "重磅", "热点"]:
            if sig in title:
                score += 2
                break

        # ═══ Title quality (0-5) ═══
        tlen = len(title)
        if 20 <= tlen <= 100:
            score += 3
        elif 15 <= tlen <= 150:
            score += 1
        if any(c.isdigit() for c in title):
            score += 2

        # ═══ Penalties ═══
        if tlen < 12:
            score -= 6
        if "?" in title or "？" in title:
            score -= 3
        cjk = sum(1 for c in title if "一" <= c <= "鿿")
        if cjk == 0:
            score -= 6
        # Hard filter: dictionary/definition results
        for bad in ["是什么意思", "有什么区别", "百度百科", "维基百科",
                     "词的区别", "区别解释",
                     "汉语词典", "新华字典", "的拼音", "组词_", "怎么读_"]:
            if bad in title:
                score -= 20
                break
        # Bing homepage/SEO spam detection: titles with many underscores (site taglines)
        # e.g. "跨境_跨境电商品牌出海企业服务平台" or "什么是跨境电商_..."
        underscore_count = title.count("_")
        if underscore_count >= 3:
            score -= 15  # keyword-stuffed SEO tagline, not an article
        elif underscore_count >= 1 and tlen < 40:
            score -= 8   # likely a site tagline, not an article title
        # Bing: "XXX | XXX - XXX" pattern often = site description, not article
        pipe_count = title.count("|")
        if pipe_count >= 3 and tlen < 60:
            score -= 10
        # Titles that are just domain names or service descriptions
        homepage_signals = ["服务平台", "一站式", "解决方案", "欢迎来到", "官网", "首页"]
        if any(s in title for s in homepage_signals) and tlen < 35:
            score -= 10

        return score

    # Attach scores to result dicts before sorting/filtering
    # Add small random jitter (±2) so similarly-scored articles shuffle between searches
    import random as _score_rng
    for r in all_results:
        r["score"] = _search_score(r) + _score_rng.randint(-2, 2)

    all_results.sort(key=lambda r: r.get("score", 0), reverse=True)
    # 动态分数门槛：结果少时放宽，确保有足够输出
    _min_score = 5 if len(all_results) < 10 else 10
    all_results = [r for r in all_results if r.get("score", 0) >= _min_score]

    # ── 质量过滤 ──
    all_results = _filter_quality_results(all_results, keyword=keyword)

    # ── 源多样性打散 ──
    # 防止同一来源（域名）霸榜：每个域名最多占 max_results 的 30%
    max_per_domain = max(2, max_results * 3 // 10)  # at least 2, at most 30%
    diverse = []
    domain_counts = {}
    _fallback = []  # 超限的文章放这里，兜底用

    def _extract_domain(url):
        """从 URL 提取域名。"""
        m = re.search(r'https?://([^/:]+)', url)
        return m.group(1).lower() if m else "unknown"

    for r in all_results:
        domain = _extract_domain(r.get("url", ""))
        cnt = domain_counts.get(domain, 0)
        if cnt < max_per_domain:
            diverse.append(r)
            domain_counts[domain] = cnt + 1
        else:
            _fallback.append(r)

    # 如果打散后不够 max_results，从 fallback 补齐
    if len(diverse) < max_results:
        for r in _fallback:
            if len(diverse) >= max_results:
                break
            diverse.append(r)

    results = diverse[:max_results]

    # ── 第 3 层：内链扩散（从 Top 微信文章发现更多相关文章）──
    if any("mp.weixin.qq.com" in r.get("url", "") for r in results[:5]):
        _internal_discovered = []
        _seen_wx_urls = set(r.get("url", "").rstrip("/") for r in results)
        for r in results[:3]:  # 只对 Top 3 做扩散，控制延迟
            url = r.get("url", "")
            if "mp.weixin.qq.com" not in url:
                continue
            try:
                links = _extract_internal_links(url, max_links=3)
                for lk in links:
                    lk_url = lk.get("url", "").rstrip("/")
                    if lk_url not in _seen_wx_urls:
                        _seen_wx_urls.add(lk_url)
                        lk["search_query"] = f"internal:{r.get('search_query', keyword)}"
                        _internal_discovered.append(lk)
            except Exception as e:
                print(f"  [InternalLink] Error: {e}")
        results.extend(_internal_discovered[:max_results])
        if _internal_discovered:
            print(f"  [InternalLink] Discovered {len(_internal_discovered)} new articles via internal links")

    # ── AI 选题增强（与 /api/news-topics 保持一致）──
    results = _enrich_news_with_topics(results, max_results)

    return jsonify({
        "keyword": keyword, "results": results, "total": len(results),
    })


@app.route("/api/news-topics", methods=["GET"])
def api_news_topics():
    """获取缓存的热点选题列表（6 小时 TTL）。"""
    limit = request.args.get("limit", 10, type=int)
    force_refresh = request.args.get("refresh", "0") == "1"
    cache_path = CACHE_DIR / "news_cache.json"

    items = []
    cache_valid = False

    if not force_refresh and cache_path.exists():
        try:
            cache_data = json.loads(cache_path.read_text(encoding="utf-8"))
            gen_at = cache_data.get("generated_at", "")
            if gen_at:
                gen_dt = datetime.fromisoformat(gen_at)
                age_hours = (datetime.now(CST) - gen_dt).total_seconds() / 3600
                if age_hours < 6:
                    items = cache_data.get("items", [])
                    cache_valid = True
                    print(f"  [Topics] 使用缓存 ({len(items)} 条, {age_hours:.1f}h)")
        except Exception as e:
            print(f"  [Topics] 缓存读取错误: {e}")

    if not cache_valid or not items:
        print("  [Topics] 拉取最新资讯...")
        items = _fetch_news_rotating(limit=min(limit * 2, 25))

        if items:
            cache_data = {
                "generated_at": datetime.now(CST).isoformat(),
                "count": len(items), "items": items,
            }
            cache_path.write_text(json.dumps(cache_data, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"  [Topics] 缓存已更新 ({len(items)} 条)")

    # AI 增强选题建议
    enriched = _enrich_news_with_topics(items[:limit], limit)

    sources = list(set(i.get("source", "") for i in enriched if i.get("source")))
    return jsonify({
        "items": enriched, "count": len(enriched),
        "updated_at": datetime.now(CST).isoformat(),
        "sources": sources, "cache_valid": cache_valid,
    })


@app.route("/api/news-topics/refresh", methods=["POST"])
def api_news_topics_refresh():
    """强制刷新热点缓存。"""
    limit = request.args.get("limit", 10, type=int)
    cache_path = CACHE_DIR / "news_cache.json"
    if cache_path.exists():
        cache_path.unlink()
    # Redirect to GET logic
    return api_news_topics()


# ── Step 2: 文章生成 ──────────────────────────────────

@app.route("/api/generate-article", methods=["POST"])
def api_generate_article():
    """基于选题生成文章。Body: {news_item, style: "b2b"|"b2c", custom_angle}

    核心改进：先从源 URL 抓取真实文章内容，再作为 AI 写作素材注入 prompt。
    这样生成的文章基于原文事实，而不是只靠摘要编造。
    """
    data = request.get_json() or {}
    news_item = data.get("news_item", {})
    style = data.get("style", "b2b")
    custom_angle = data.get("custom_angle", "")
    skip_fetch = data.get("skip_fetch", False)  # 允许前端跳过抓取以加速

    if not news_item.get("title"):
        return jsonify({"error": "news_item.title 必填"}), 400

    # ── Step 0: Fetch real article content from source URL ──
    article_content = ""
    resolved_url = news_item.get("url", "")
    if not skip_fetch and resolved_url:
        print(f"  [Generate] Fetching article content from: {resolved_url[:80]}")
        article_content, resolved_url = _fetch_general_article(resolved_url, timeout=12)
        if article_content:
            news_item["_resolved_url"] = resolved_url
            print(f"  [Generate] ✓ Got {len(article_content)} chars of article content")
        else:
            print(f"  [Generate] ✗ Could not fetch article content, will use snippet fallback")

    if style == "b2p":
        system_prompt = B2P_SYSTEM_PROMPT
    elif style == "b2b":
        system_prompt = B2B_SYSTEM_PROMPT
    else:
        system_prompt = B2C_SYSTEM_PROMPT

    user_prompt = _build_article_prompt(news_item, style, custom_angle, article_content)

    print(f"  [Generate] style={style}, topic={news_item.get('suggested_topic', news_item.get('title'))[:50]}")

    content = llm_chat_text(system=system_prompt, user=user_prompt, temperature=0.7, max_tokens=3500)

    if not content:
        # Fallback
        content = _generate_fallback(news_item, style)

    # Save to output
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r"[^\w一-鿿]+", "_", news_item.get("title", "article"))[:40]
    filename = f"{timestamp}_{safe_title}.md"
    filepath = OUTPUT_DIR / "articles" / filename
    filepath.write_text(content, encoding="utf-8")

    return jsonify({
        "content": content, "style": style,
        "filename": filename, "saved_path": str(filepath),
        "char_count": len(content),
        "has_source_content": bool(article_content),
        "source_content_chars": len(article_content),
    })


def _generate_fallback(news_item, style):
    """API 不可用时的降级文章。"""
    topic = news_item.get("suggested_topic", news_item.get("title", "跨境电商"))
    if style == "b2b":
        return f"""**核心结论**: 针对「{topic}」涉及的行业趋势，跨境电商行业正处于快速发展期，卖家对合规化、专业化服务的需求日益增长。

## 行业背景

2026 年跨境电商市场规模持续扩大，但合规门槛也在同步提高。从欧盟 VAT 新政到亚马逊气候绿标要求，卖家面临多重挑战。

## 方案对比

| 维度 | 全链路服务商 | 单点代理商 | 自助 DIY |
|------|------------|-----------|---------|
| 📋 覆盖范围 | 知识产权/税务/合规/物流全链路 | 单点服务 | 自行研究 |
| 🛡️ 安全合规 | 平台官方合作资质 | 资质参差不齐 | 完全自担风险 |
| ⚡ 响应速度 | 1v1顾问，2h响应 | 3-5个工作日 | 取决于自身能力 |
| 💰 综合成本 | 按需付费，透明 | 可能有隐形费用 | 时间成本高 |

> 选择哪种方案取决于卖家规模、预算和内部团队能力。建议先梳理自身需求，再做匹配。
"""
    else:
        return f"""说实话，看到「{topic}」这个话题我太有共鸣了！😭

作为一个去年刚入坑跨境的个人卖家，我踩过的坑可能比你看过的攻略还多。

## 🔥 我的经验总结

### 核心教训：不要只看价格

我先后试过 3 种不同的方案：
- ✅ **全托管服务**: 省心但贵，适合月销 €10000+ 的卖家
- ✅ **半自助 + 顾问**: 性价比最高，自己处理简单的，难的找专业的人
- ✅ **纯 DIY**: 最省钱但最耗时，适合刚起步测款的阶段

讲真，月销 €3000 以下的微型卖家可能更适合先走半自助路线——把 VAT 申报这种容易出错的部分外包，商标注册这种一次性的事情可以自己研究。但如果月销过万欧，专业的合规服务绝对是值得的投资。

选服务最重要的是看对方是否有平台官方合作资质，以及合同条款是否透明。建议大家多对比几家再决定～

有问题评论区喊我，看到就回～"""


# ── Step 3: 图片生成 ──────────────────────────────────

@app.route("/api/generate-cover", methods=["POST"])
def api_generate_cover():
    """生成封面图 (2.35:1 横版)。Body: {article_content, custom_prompt?}"""
    data = request.get_json() or {}
    article_content = data.get("article_content", "")
    # custom_prompt 可能是 string 或 dict（前端编辑器可能传对象），兼容处理
    raw_custom = data.get("custom_prompt", "") or ""
    if isinstance(raw_custom, dict):
        raw_custom = raw_custom.get("prompt", raw_custom.get("text", str(raw_custom)))
    custom_prompt = (raw_custom or "").strip() if isinstance(raw_custom, str) else ""

    if not article_content and not custom_prompt:
        return jsonify({"error": "article_content 或 custom_prompt 必填"}), 400

    if custom_prompt:
        prompt = custom_prompt
        print(f"  [Cover] Custom prompt: {prompt[:80]}...")
    else:
        prompt = _build_image_prompt(article_content, "cover")
        print(f"  [Cover] Prompt: {prompt[:80]}...")

    urls = _generate_ark_image(prompt, size="cover", n=1)

    if not urls:
        return jsonify({
            "error": "封面生成失败：火山方舟 API 未返回图片，可能是 prompt 触发了内容审核或 API 暂时不可用，请稍后重试",
            "prompt": prompt, "image_urls": [], "image_type": "cover",
        }), 500

    # Download locally
    saved = []
    for i, url in enumerate(urls):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_path = OUTPUT_DIR / "images" / f"cover_{timestamp}_{i}.png"
        if _download_image(url, save_path):
            saved.append(str(save_path))

    return jsonify({
        "prompt": prompt, "image_urls": urls,
        "saved_paths": saved, "image_type": "cover",
    })


@app.route("/api/generate-image", methods=["POST"])
def api_generate_image():
    """生成正文配图。Body: {article_content, count: 3, custom_prompts?: ["prompt1",...]}"""
    data = request.get_json() or {}
    article_content = data.get("article_content", "")
    count = min(int(data.get("count", 3)), 4)
    custom_prompts = data.get("custom_prompts", None)

    if not article_content and not custom_prompts:
        return jsonify({"error": "article_content 或 custom_prompts 必填"}), 400

    # If custom prompts provided, use them directly
    if custom_prompts and isinstance(custom_prompts, list) and len(custom_prompts) > 0:
        results = []
        for i, cp in enumerate(custom_prompts[:count]):
            cp = (cp or "").strip()
            if not cp:
                continue
            print(f"  [Image {i+1}] Custom prompt: {cp[:80]}...")
            urls = _generate_ark_image(cp, size="body", n=1)
            saved = []
            for j, url in enumerate(urls):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                save_path = OUTPUT_DIR / "images" / f"body_{timestamp}_{i}_{j}.png"
                if _download_image(url, save_path):
                    saved.append(str(save_path))
            results.append({
                "index": i, "prompt": cp, "image_urls": urls,
                "saved_paths": saved, "section_excerpt": cp[:100],
            })
        return jsonify({"images": results, "total": len(results)})

    # Otherwise extract from article content
    paragraphs = [p.strip() for p in article_content.split("\n\n") if len(p.strip()) > 50]
    key_sections = paragraphs[:count] if len(paragraphs) >= count else paragraphs

    results = []
    for i, section in enumerate(key_sections):
        prompt = _build_image_prompt(section[:200], "body")
        print(f"  [Image {i+1}] Prompt: {prompt[:80]}...")

        urls = _generate_ark_image(prompt, size="body", n=1)

        saved = []
        for j, url in enumerate(urls):
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            save_path = OUTPUT_DIR / "images" / f"body_{timestamp}_{i}_{j}.png"
            if _download_image(url, save_path):
                saved.append(str(save_path))

        results.append({
            "index": i, "prompt": prompt, "image_urls": urls,
            "saved_paths": saved, "section_excerpt": section[:100],
        })

    return jsonify({"images": results, "total": len(results)})


# ── WeChat HTML Converter ──────────────────────────────

# Theme presets for WeChat article formatting
_WECHAT_THEMES = {
    "default": {
        "name": "经典红金",
        "bg": "#ffffff",
        "text": "#3e3e3e",
        "h1_color": "#1a1a1a", "h1_accent": "#CE0E19",
        "h2_color": "#2c2c2c", "h3_color": "#3a3a3a",
        "strong_color": "#CE0E19",
        "quote_border": "#D9A85A", "quote_bg": "#fef9e7", "quote_text": "#5a4a2a",
        "table_header_bg": "#f5f5f5",
        "link_color": "#CE0E19",
    },
    "warm": {
        "name": "暖橙生活",
        "bg": "#ffffff",
        "text": "#4a3f35",
        "h1_color": "#3d2e1e", "h1_accent": "#E8833A",
        "h2_color": "#5c4033", "h3_color": "#6b5344",
        "strong_color": "#E8833A",
        "quote_border": "#E8833A", "quote_bg": "#fef8f3", "quote_text": "#6b4c3b",
        "table_header_bg": "#fdf3ea",
        "link_color": "#E8833A",
    },
    "business": {
        "name": "深蓝商务",
        "bg": "#ffffff",
        "text": "#2c3e50",
        "h1_color": "#1a2a3a", "h1_accent": "#1a5276",
        "h2_color": "#1a3a4a", "h3_color": "#2c3e50",
        "strong_color": "#1a5276",
        "quote_border": "#2980b9", "quote_bg": "#eaf2f8", "quote_text": "#1a5276",
        "table_header_bg": "#eaf2f8",
        "link_color": "#1a5276",
    },
    "minimal": {
        "name": "极简黑白",
        "bg": "#ffffff",
        "text": "#333333",
        "h1_color": "#000000", "h1_accent": "#333333",
        "h2_color": "#1a1a1a", "h3_color": "#333333",
        "strong_color": "#000000",
        "quote_border": "#999999", "quote_bg": "#f8f8f8", "quote_text": "#555555",
        "table_header_bg": "#f0f0f0",
        "link_color": "#333333",
    },
    "elegant": {
        "name": "雅致金棕",
        "bg": "#ffffff",
        "text": "#4a3f35",
        "h1_color": "#2d1f0e", "h1_accent": "#B8860B",
        "h2_color": "#3d2e1e", "h3_color": "#5c4033",
        "strong_color": "#B8860B",
        "quote_border": "#B8860B", "quote_bg": "#fdf8ee", "quote_text": "#6b4c3b",
        "table_header_bg": "#fdf5e1",
        "link_color": "#B8860B",
    },
}


def _wechat_style_block(tag, styles):
    """Build an inline style string from a dict."""
    return f'{tag} style="{"; ".join(f"{k}:{v}" for k,v in styles.items())}"'


def _markdown_to_wechat_html(md_text, title="", theme="default", font_size=15,
                              accent_color=None, images=None):
    """将 Markdown 转换为微信公众号兼容的 HTML（全内联样式）。

    Args:
        md_text: Markdown source text
        title: Article title (used in <title>)
        theme: Theme preset name (default/warm/business/minimal/elegant)
        font_size: Base font size in px (14-18)
        accent_color: Override accent/strong color (hex, e.g. "#CE0E19")
        images: Optional dict mapping image index to URL for embedding
    """
    import re as _re

    t = _WECHAT_THEMES.get(theme, _WECHAT_THEMES["default"]).copy()
    if accent_color:
        t["strong_color"] = accent_color
        t["h1_accent"] = accent_color
        t["link_color"] = accent_color

    fs = max(14, min(22, font_size))
    fs_h1 = fs + 7
    fs_h2 = fs + 3
    fs_h3 = fs + 1
    fs_small = fs - 1

    # ── Phase 1: Line-by-line processing with state ──
    lines = md_text.split('\n')
    out_lines = []
    in_blockquote = False
    bq_lines = []
    in_list = False
    list_lines = []
    list_type = None
    in_table = False
    table_rows = []

    def flush_blockquote():
        nonlocal in_blockquote, bq_lines
        if not bq_lines:
            return
        content = '<br>'.join(bq_lines)
        html = (
            f'<blockquote style="border-left:3px solid {t["quote_border"]};'
            f'padding:8px 12px;margin:14px 0;background:{t["quote_bg"]};'
            f'color:{t["quote_text"]};font-size:{fs_small}px;'
            f'border-radius:0 4px 4px 0;line-height:1.8">'
            f'{content}</blockquote>'
        )
        out_lines.append(html)
        bq_lines = []
        in_blockquote = False

    def flush_list():
        nonlocal in_list, list_lines, list_type
        if not list_lines:
            return
        tag = 'ol' if list_type == 'ordered' else 'ul'
        items = ''.join(
            f'<li style="font-size:{fs}px;color:{t["text"]};line-height:1.8;margin-bottom:6px">{li}</li>'
            for li in list_lines
        )
        html = f'<{tag} style="padding-left:20px;margin:10px 0">{items}</{tag}>'
        out_lines.append(html)
        list_lines = []
        in_list = False
        list_type = None

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        html = (
            f'<table style="border-collapse:collapse;width:100%;'
            f'margin:14px 0;font-size:{fs_small}px"><tbody>'
        )
        for ri, row in enumerate(table_rows):
            tag = 'th' if ri == 0 else 'td'
            bg = (f' style="border:1px solid #e0e0e0;padding:8px 10px;'
                  f'text-align:left;background:{t["table_header_bg"]};'
                  f'font-weight:600;color:#333"'
                  if ri == 0 else
                  f' style="border:1px solid #e0e0e0;padding:8px 10px;text-align:left"')
            cells_html = ''.join(f'<{tag}{bg}>{c}</{tag}>' for c in row)
            html += f'<tr>{cells_html}</tr>'
        html += '</tbody></table>'
        out_lines.append(html)
        table_rows = []
        in_table = False

    def process_inline(text):
        """Process inline formatting: bold, code, links."""
        # Bold **...**
        text = _re.sub(
            r'\*\*(.+?)\*\*',
            rf'<strong style="color:{t["strong_color"]};font-weight:600">\1</strong>',
            text
        )
        # Inline code `...`
        text = _re.sub(
            r'`([^`]+)`',
            rf'<code style="background:#f1f5f9;padding:2px 6px;border-radius:3px;font-size:90%;color:{t["strong_color"]}">\1</code>',
            text
        )
        # Links [text](url)
        text = _re.sub(
            r'\[([^\]]+)\]\(([^)]+)\)',
            rf'<a style="color:{t["link_color"]};text-decoration:none;border-bottom:1px solid {t["link_color"]}" href="\2">\1</a>',
            text
        )
        return text

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Empty line: flush all states ──
        if not line.strip():
            if in_blockquote:
                flush_blockquote()
            if in_list:
                flush_list()
            if in_table:
                flush_table()
            out_lines.append('')
            i += 1
            continue

        # ── Table ──
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                flush_blockquote()
                flush_list()
                in_table = True
                table_rows = []
            # Skip separator
            if _re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            cells = [process_inline(c.strip()) for c in line.strip().split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue

        # ── Blockquote ──
        if line.startswith('>'):
            if in_list:
                flush_list()
            if in_table:
                flush_table()
            if not in_blockquote:
                in_blockquote = True
                bq_lines = []
            # Remove leading > and one optional space
            content = _re.sub(r'^>\s?', '', line)
            bq_lines.append(process_inline(content))
            i += 1
            continue

        # ── Unordered list ──
        list_match = _re.match(r'^([\-*])\s+(.+)$', line)
        if list_match:
            if in_blockquote:
                flush_blockquote()
            if in_table:
                flush_table()
            if not in_list or list_type != 'unordered':
                flush_list()
                in_list = True
                list_type = 'unordered'
            list_lines.append(process_inline(list_match.group(2)))
            i += 1
            continue

        # ── Ordered list ──
        ol_match = _re.match(r'^(\d+)[\.\)]\s+(.+)$', line)
        if ol_match:
            if in_blockquote:
                flush_blockquote()
            if in_table:
                flush_table()
            if not in_list or list_type != 'ordered':
                flush_list()
                in_list = True
                list_type = 'ordered'
            list_lines.append(process_inline(ol_match.group(2)))
            i += 1
            continue

        # ── Flush any open state on non-matching line ──
        if in_blockquote:
            flush_blockquote()
        if in_list:
            flush_list()
        if in_table:
            flush_table()

        # ── Heading ──
        h1 = _re.match(r'^#\s+(.+)$', line)
        h2 = _re.match(r'^##\s+(.+)$', line)
        h3 = _re.match(r'^###\s+(.+)$', line)
        h4 = _re.match(r'^####\s+(.+)$', line)

        if h1:
            content = process_inline(h1.group(1))
            html = (
                f'<h1 style="font-size:{fs_h1}px;font-weight:700;'
                f'color:{t["h1_color"]};margin:18px 0 14px;'
                f'padding-left:12px;border-left:4px solid {t["h1_accent"]};'
                f'line-height:1.4">{content}</h1>'
            )
            out_lines.append(html)
        elif h2:
            content = process_inline(h2.group(1))
            html = (
                f'<h2 style="font-size:{fs_h2}px;font-weight:700;'
                f'color:{t["h2_color"]};margin:16px 0 10px;line-height:1.4;'
                f'padding-bottom:8px;border-bottom:2px solid {t["strong_color"]}">'
                f'{content}</h2>'
            )
            out_lines.append(html)
        elif h3:
            content = process_inline(h3.group(1))
            html = (
                f'<h3 style="font-size:{fs_h3}px;font-weight:600;'
                f'color:{t["h3_color"]};margin:14px 0 8px;line-height:1.4">'
                f'{content}</h3>'
            )
            out_lines.append(html)
        elif h4:
            content = process_inline(h4.group(1))
            html = (
                f'<h4 style="font-size:{fs}px;font-weight:600;'
                f'color:#4a4a4a;margin:12px 0 6px;line-height:1.4">'
                f'{content}</h4>'
            )
            out_lines.append(html)
        # ── Horizontal rule ──
        elif _re.match(r'^[-*_]{3,}$', line.strip()):
            out_lines.append('<hr style="border:none;border-top:1px solid #e8e8e8;margin:20px 0">')
        # ── Regular paragraph ──
        else:
            content = process_inline(line)
            html = (
                f'<p style="font-size:{fs}px;color:{t["text"]};'
                f'line-height:1.85;margin:0 0 14px;letter-spacing:.5px;'
                f'text-align:justify">{content}</p>'
            )
            out_lines.append(html)

        i += 1

    # Flush any remaining state at end
    flush_blockquote()
    flush_list()
    flush_table()

    # ── Phase 2: Insert images at strategic positions ──
    body_html = '\n'.join(out_lines)
    if images and isinstance(images, dict) and len(images) > 0:
        # Sort: "cover" first, then numeric keys
        def _img_sort_key(item):
            k = item[0]
            if k == "cover":
                return (-1, "")
            try:
                return (int(k), "")
            except (ValueError, TypeError):
                return (9999, k)
        img_urls = [v for k, v in sorted(images.items(), key=_img_sort_key) if v]
        if img_urls:
            # Find all paragraph-ending positions
            end_positions = [m.end() for m in _re.finditer(
                r'(</p>|</h[1-4]>|</blockquote>|</table>|</ul>|</ol>|</li>|</hr>)', body_html
            )]
            if len(end_positions) >= 2 and len(img_urls) >= 2:
                # Cover after first block element
                insert_pos = end_positions[0]
                cover_html = (
                    f'<p style="text-align:center;margin:16px 0">'
                    f'<img src="{img_urls[0]}" '
                    f'style="max-width:100%;height:auto;border-radius:6px;box-shadow:0 2px 8px rgba(0,0,0,.1)" '
                    f'alt="封面图"></p>'
                )
                body_html = body_html[:insert_pos] + '\n' + cover_html + '\n' + body_html[insert_pos:]

                # Distribute body images across remaining paragraphs
                body_imgs = img_urls[1:]
                if body_imgs and len(end_positions) > 3:
                    step = max(1, (len(end_positions) - 1) // (len(body_imgs) + 1))
                    accumulated_shift = len(cover_html) + 2
                    for idx, img_url in enumerate(body_imgs):
                        split_idx = min((idx + 1) * step + 1, len(end_positions) - 1)
                        pos = end_positions[split_idx] + accumulated_shift
                        img_html = (
                            f'<p style="text-align:center;margin:16px 0">'
                            f'<img src="{img_url}" '
                            f'style="max-width:100%;height:auto;border-radius:4px" '
                            f'alt="配图{idx + 1}"></p>'
                        )
                        body_html = body_html[:pos] + '\n' + img_html + '\n' + body_html[pos:]
                        accumulated_shift += len(img_html) + 2
            elif len(img_urls) >= 1:
                # Just one image: insert after first paragraph
                first_p = body_html.find('</p>')
                if first_p > 0:
                    img_html = (
                        f'<p style="text-align:center;margin:16px 0">'
                        f'<img src="{img_urls[0]}" '
                        f'style="max-width:100%;height:auto;border-radius:6px" '
                        f'alt="配图"></p>'
                    )
                    body_html = body_html[:first_p + 4] + '\n' + img_html + '\n' + body_html[first_p + 4:]

    # ── Phase 3: Wrap in full document ──
    full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1,user-scalable=no">
<title>{title or '微信公众号文章'}</title>
</head>
<body style="max-width:677px;margin:0 auto;padding:16px;background:{t['bg']};font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif">
{body_html}
</body>
</html>"""

    return full_html


@app.route("/api/to-wechat-html", methods=["POST"])
def api_to_wechat_html():
    """将 Markdown 文章转换为微信公众号兼容的 HTML。
    Body: {
        content: "...", title: "...",
        theme: "default"|"warm"|"business"|"minimal"|"elegant",
        font_size: 15, accent_color: "#CE0E19",
        images: {"cover": "url", "0": "url", ...}
    }
    Returns: {html: "...", char_count: N, theme: "..."}
    """
    data = request.get_json() or {}
    content = data.get("content", "")
    title = data.get("title", "")
    theme = data.get("theme", "default")
    font_size = int(data.get("font_size", 15))
    accent_color = data.get("accent_color", None)
    images = data.get("images", None)

    if not content:
        return jsonify({"error": "content 必填"}), 400

    html = _markdown_to_wechat_html(
        content, title, theme=theme, font_size=font_size,
        accent_color=accent_color, images=images
    )

    return jsonify({
        "html": html,
        "char_count": len(html),
        "theme": theme,
        "font_size": font_size,
        "wechat_compatible": True,
        "themes_available": list(_WECHAT_THEMES.keys()),
    })


# ── Article Rewrite ────────────────────────────────────

def _fetch_article_content(url):
    """Universal article content extractor — works across platforms.

    Extraction strategies (tried in order):
      1. trafilatura — handles 90% of news/article sites
      2. BeautifulSoup with generic selectors — fallback
      3. Jina Reader API — last resort (handles JS-rendered pages)

    Args:
        url: Any article/news URL

    Returns:
        dict with keys: title, content, author, source_url, extraction_method,
        or None on failure
    """
    import re as _re
    import urllib.request as _ur

    url = _resolve_wechat_url(url, timeout=10)

    USER_AGENT = (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )

    html = None

    # ── Strategy 1: Direct fetch + trafilatura ──
    try:
        req = _ur.Request(url, headers={"User-Agent": USER_AGENT})
        with _ur.urlopen(req, timeout=15) as resp:
            html = resp.read()
        # trafilatura handles encoding detection internally
        import trafilatura as _traf
        result = _traf.extract(
            html,
            include_comments=False,
            include_tables=False,
            output_format="markdown",
            url=url,
            with_metadata=True,
        )
        if result and len(result.strip()) >= 80:
            # Anti-bot guard: check for captcha/verify page signals
            _anti_bot_kw = ["环境异常", "验证", "captcha", "验证码", "请确认",
                            "安全验证", "请输入", "访问受限", "js_verify",
                            "are you a robot", "请点击", "security check"]
            _result_lower = result[:200].lower()
            _is_anti_bot = any(kw in result[:500] or kw.lower() in _result_lower
                               for kw in _anti_bot_kw)
            if _is_anti_bot:
                print(f"  [Fetch] trafilatura returned anti-bot page, falling back...")

            if not _is_anti_bot and len(result.strip()) >= 200:
                # Try to get metadata
                meta = _traf.extract(html, output_format="json", url=url, with_metadata=True)
                title = ""
                author = ""
                if meta:
                    try:
                        import json as _json
                        meta_obj = _json.loads(meta)
                        title = (meta_obj.get("title") or "").strip()
                        author = (meta_obj.get("author") or "").strip()
                    except Exception:
                        pass

                # If trafilatura didn't get title, try HTML fallback
                if not title:
                    text = html.decode("utf-8", errors="replace")
                    og_m = _re.search(r'<meta[^>]+property="og:title"[^>]+content="([^"]+)"', text)
                    if og_m:
                        title = og_m.group(1).strip()
                    if not title:
                        title_m = _re.search(r"<title>([^<]+)</title>", text)
                        if title_m:
                            title = title_m.group(1).strip()

                if not _core_extractors.is_usable_article_content(result, url):
                    print(f"  [Fetch] trafilatura content incomplete after cleanup, falling back...")
                else:
                    print(f"  [Fetch] trafilatura OK: {len(result)} chars, title={title[:40]}")
                    return _core_extractors.article_success(
                        title=title or "未知标题",
                        author=author,
                        content=result,
                        source_url=url,
                        extraction_method="trafilatura",
                    )
    except Exception as e:
        print(f"  [Fetch] trafilatura failed: {e}")

    # ── Strategy 1.5: curl_cffi (browser TLS fingerprint impersonation) ──
    try:
        from curl_cffi import requests as _cc_req
        _cc_resp = _cc_req.get(url, impersonate="chrome120", timeout=15,
                               headers={"Accept-Language": "zh-CN,zh;q=0.9",
                                        "Accept": "text/html,application/xhtml+xml",
                                        "User-Agent": USER_AGENT})
        _cc_html = _cc_resp.text
        # Anti-bot guard
        _cc_lower = _cc_html[:500].lower()
        _is_anti_bot = any(
            kw in _cc_html[:500] or kw.lower() in _cc_lower
            for kw in ["环境异常", "验证码", "captcha", "js_verify", "安全验证"]
        )
        if not _is_anti_bot and len(_cc_html) > 5000:
            # Got real article HTML — extract with trafilatura
            import trafilatura as _traf2
            _cc_content = _traf2.extract(
                _cc_html, include_comments=False, include_tables=False,
                output_format="markdown", url=url, with_metadata=True
            )
            if _cc_content and len(_cc_content.strip()) >= 200:
                _cc_title = ""
                _cc_author = ""
                try:
                    _cc_meta = _traf2.extract(_cc_html, output_format="json", url=url, with_metadata=True)
                    if _cc_meta:
                        import json as _json2
                        _cc_meta_obj = _json2.loads(_cc_meta)
                        _cc_title = (_cc_meta_obj.get("title") or "").strip()
                        _cc_author = (_cc_meta_obj.get("author") or "").strip()
                except Exception:
                    pass
                if not _cc_title:
                    _og_m = _re.search(r'<meta[^>]+property="og:title"[^>]+content="([^"]+)"', _cc_html)
                    if _og_m:
                        _cc_title = _og_m.group(1).strip()
                    if not _cc_title:
                        _t_m = _re.search(r"<title>([^<]+)</title>", _cc_html)
                        if _t_m:
                            _cc_title = _t_m.group(1).strip()
                if not _core_extractors.is_usable_article_content(_cc_content, url):
                    print(f"  [Fetch] curl_cffi content incomplete after cleanup, falling back...")
                else:
                    print(f"  [Fetch] curl_cffi OK: {len(_cc_content)} chars, title={_cc_title[:40]}")
                    return _core_extractors.article_success(
                        title=_cc_title or "未知标题",
                        author=_cc_author,
                        content=_cc_content,
                        source_url=url,
                        extraction_method="curl_cffi",
                    )
            else:
                print(f"  [Fetch] curl_cffi got HTML but trafilatura extraction short")
        else:
            print(f"  [Fetch] curl_cffi got anti-bot page ({len(_cc_html)} bytes), falling back...")
    except ImportError:
        pass  # curl_cffi not installed
    except Exception as e:
        print(f"  [Fetch] curl_cffi failed: {e}")

    # ── Strategy 2: BeautifulSoup with generic selectors ──
    try:
        from bs4 import BeautifulSoup as _BS

        if html is None:
            req = _ur.Request(url, headers={"User-Agent": USER_AGENT})
            with _ur.urlopen(req, timeout=15) as resp:
                html = resp.read()

        text = html.decode("utf-8", errors="replace")
        soup = _BS(text, "html.parser")

        # Extract title
        title = ""
        for sel in [
            ("meta", {"property": "og:title"}),
            ("meta", {"name": "twitter:title"}),
            ("h1", {"class": _re.compile(r"title|article|post|entry", re.I)}),
            ("h1", {}),
            ("title", {}),
        ]:
            if isinstance(sel[1], dict) and "class" in sel[1] and hasattr(sel[1]["class"], "search"):
                # Cannot use compiled regex with find() directly
                found = soup.find("h1", class_=sel[1]["class"])
                if found:
                    title = found.get_text(strip=True)
                    break
            else:
                found = soup.find(sel[0], sel[1])
                if found:
                    title = (found.get("content", "") or found.get_text(strip=True))
                    break
        title = title.strip()

        # Extract author
        author = ""
        for meta_name in ["author", "article:author", "og:article:author"]:
            author_meta = soup.find("meta", attrs={"name": meta_name})
            if not author_meta:
                author_meta = soup.find("meta", property=meta_name)
            if author_meta and author_meta.get("content"):
                author = author_meta["content"].strip()
                break

        # Extract content — try common article containers
        content_div = None
        for sel in [
            ("div", {"id": "js_content"}),           # WeChat
            ("div", {"class": "rich_media_content"}), # WeChat v2
            ("article", {}),                          # HTML5 article tag
            ("div", {"id": _re.compile(r"content|article|post|entry", re.I)}),
            ("div", {"class": _re.compile(r"content|article|post|entry|body|text", re.I)}),
            ("main", {}),                             # HTML5 main tag
            ("section", {"class": _re.compile(r"article|content", re.I)}),
        ]:
            if isinstance(sel[1], dict) and any(isinstance(v, type(_re.compile(""))) for v in sel[1].values()
                                                 if hasattr(v, "search")):
                # Skip regex-based attrs for find() — handle separately
                continue
            content_div = soup.find(sel[0], sel[1])
            if content_div:
                break

        # Fallback: try article > section > div that has most text
        if not content_div:
            candidates = soup.find_all(["article", "section", "div"])
            best = None
            best_len = 0
            for c in candidates:
                # Skip nav, header, footer, sidebar
                if c.get("id") and any(kw in str(c["id"]).lower()
                                       for kw in ("nav", "header", "footer", "sidebar", "menu", "comment")):
                    continue
                if c.get("class") and any(kw in " ".join(c.get("class", [])).lower()
                                          for kw in ("nav", "header", "footer", "sidebar", "menu", "comment")):
                    continue
                text_len = len(c.get_text(strip=True))
                if text_len > best_len and text_len > 200:
                    best = c
                    best_len = text_len
            content_div = best

        if content_div:
            for tag in content_div.find_all(["script", "style", "nav", "footer"]):
                tag.decompose()
            content = content_div.get_text(separator="\n", strip=True)
        else:
            # Last resort: body text
            body = soup.find("body")
            if body:
                for tag in body.find_all(["script", "style", "nav", "footer", "header"]):
                    tag.decompose()
                content = body.get_text(separator="\n", strip=True)
            else:
                print("  [Fetch] No content found in page")
                return None

        content = _re.sub(r"\n\s*\n\s*\n+", "\n\n", content).strip()

        # Anti-bot guard for BeautifulSoup extraction too
        _anti_bot_kw = ["环境异常", "验证码", "captcha", "安全验证", "请输入验证",
                        "访问受限", "js_verify", "are you a robot", "请点击"]
        _is_anti_bot = any(kw in content[:500] or kw.lower() in content[:500].lower()
                           for kw in _anti_bot_kw)

        if _is_anti_bot:
            print(f"  [Fetch] BeautifulSoup returned anti-bot page, falling back...")
        elif len(content) < 80:
            print(f"  [Fetch] BeautifulSoup content too short: {len(content)} chars")
        else:
            print(f"  [Fetch] BeautifulSoup OK: {len(content)} chars, title={title[:40]}")
            return _core_extractors.article_success(
                title=title or "未知标题",
                author=author,
                content=content,
                source_url=url,
                extraction_method="beautifulsoup",
            )
    except Exception as e:
        print(f"  [Fetch] BeautifulSoup failed: {e}")

    # ── Strategy 3: Jina Reader API ──
    try:
        jina_url = f"https://r.jina.ai/{url}"
        req = _ur.Request(jina_url, headers={
            "Accept": "text/markdown",
            "User-Agent": USER_AGENT,
        })
        with _ur.urlopen(req, timeout=20) as resp:
            md_text = resp.read().decode("utf-8", errors="replace")

        if md_text and len(md_text) >= 80:
            # Jina prepends title as markdown header
            title = ""
            lines = md_text.split("\n")
            if lines and lines[0].startswith("# "):
                title = lines[0][2:].strip()
                # Remove the jina origin line if present
                content_start = 1
                while content_start < len(lines) and (
                    lines[content_start].startswith(">") and "jina.ai" in lines[content_start]
                    or not lines[content_start].strip()
                ):
                    content_start += 1
                content = "\n".join(lines[content_start:]).strip()
            else:
                content = md_text.strip()

            print(f"  [Fetch] Jina Reader OK: {len(content)} chars")
            return _core_extractors.article_success(
                title=title or "未知标题",
                author="",
                content=content,
                source_url=url,
                extraction_method="jina_reader",
            )
    except Exception as e:
        print(f"  [Fetch] Jina Reader failed: {e}")

    print("  [Fetch] All strategies exhausted — cannot extract")
    return None


@app.route("/api/fetch-article", methods=["POST"])
def api_fetch_article():
    """Fetch any article by URL and return its content.

    Supports: WeChat, 知乎, 36氪, 亿邦动力, 搜狐, 网易, and any standard article page.

    Body: {url: "https://..."}
    Returns: {title, author, content, char_count, source_url, extraction_method}
    """
    data = request.get_json() or {}
    url = (data.get("url") or "").strip()

    if not url:
        return jsonify({"error": "url 必填"}), 400

    article = _fetch_article_content(url)
    if not article:
        failure = _core_extractors.article_failure(
            url,
            status="extract_failed",
            error_hint="服务器端提取失败（可能被反爬拦截），建议在浏览器中打开文章后复制全文导入",
        )
        return jsonify({
            "error": "无法提取文章内容，请检查链接是否正确，或尝试直接粘贴文章全文",
            "hint": failure["error_hint"],
            "status": failure["status"],
            "error_hint": failure["error_hint"],
            "manual_import_recommended": failure["manual_import_recommended"],
            "source_url": failure["source_url"],
            "extraction_method": failure["extraction_method"],
        }), 422

    return jsonify(_core_extractors.normalize_article_result(article, fallback_url=url))


def _summarize_article_check(url, article):
    """Return a metadata-only extraction summary for diagnostics."""
    if not article:
        failure = _core_extractors.article_failure(url)
        return {
            "url": url,
            "ok": False,
            "status": failure["status"],
            "method": failure["extraction_method"],
            "title": "",
            "chars": 0,
            "quality": failure["quality"],
            "manual_import_recommended": failure["manual_import_recommended"],
        }

    normalized = _core_extractors.normalize_article_result(article, fallback_url=url)
    return {
        "url": url,
        "ok": normalized["ok"],
        "status": normalized["status"],
        "method": normalized["extraction_method"],
        "title": normalized["title"],
        "chars": normalized["char_count"],
        "quality": normalized["quality"],
        "manual_import_recommended": normalized.get("manual_import_recommended", False),
    }


@app.route("/api/verify-wechat-links", methods=["POST"])
def api_verify_wechat_links():
    """Batch-check article extraction quality without returning article bodies."""
    data = request.get_json() or {}
    urls = data.get("urls") or []
    if isinstance(urls, str):
        urls = [urls]
    urls = [u.strip() for u in urls if isinstance(u, str) and u.strip()]
    if not urls:
        return jsonify({"error": "urls 必填"}), 400
    if len(urls) > 20:
        return jsonify({"error": "一次最多验证 20 个链接"}), 400

    results = []
    for url in urls:
        try:
            article = _fetch_article_content(url)
            results.append(_summarize_article_check(url, article))
        except Exception as e:
            failure = _core_extractors.article_failure(url, error_hint=f"{type(e).__name__}: {e}")
            results.append(_summarize_article_check(url, failure))

    usable_count = sum(1 for r in results if r.get("ok") and (r.get("quality") or {}).get("usable"))
    return jsonify({
        "count": len(results),
        "usable_count": usable_count,
        "results": results,
    })


@app.route("/api/rewrite", methods=["POST"])
def api_rewrite_article():
    """Rewrite any article in the specified B2P/B2B/B2C style.

    Body: {
        url: "https://...",                      # 文章链接 (任意平台)
        raw_content: "...",                       # 直接输入的文章内容 (与url二选一)
        style: "b2p" | "b2b" | "b2c",            # 写作风格，默认 b2p
        custom_angle: "...",                      # 可选: 自定义写作角度
        theme: "default" | "warm" | ...           # WeChat HTML 主题，默认 default
    }
    Returns: {
        original_title, original_author, original_char_count,
        rewritten_markdown, rewritten_html, style, theme,
        extraction_method, wechat_compatible
    }
    """
    data = request.get_json() or {}

    url = (data.get("url") or "").strip()
    raw_content = (data.get("raw_content") or "").strip()
    style = (data.get("style") or "b2p").strip().lower()
    custom_angle = (data.get("custom_angle") or "").strip()
    theme = (data.get("theme") or "default").strip()
    use_raw_content = bool(raw_content and len(raw_content) >= MIN_RAW_REWRITE_CHARS)
    import_info = _classify_article_import(url if not use_raw_content else "", raw_content if use_raw_content else "")

    if style not in ("b2p", "b2b", "b2c"):
        return jsonify({"error": "style 必须是 b2p, b2b 或 b2c"}), 400

    extraction_method = "raw_input"

    # ── Get source content ──
    source_url = ""
    source_quality = _core_extractors.assess_article_quality(raw_content, source_url="")
    if raw_content and not use_raw_content and not url:
        return jsonify({
            "error": f"原文内容太短，至少需要 {MIN_RAW_REWRITE_CHARS} 字；请粘贴完整文章或提供文章链接",
            "status": "raw_content_too_short",
            "source_quality": source_quality,
        }), 400

    if use_raw_content:
        original_title = data.get("original_title", "直接输入内容")
        original_author = ""
        source_content = raw_content
        source_quality = _core_extractors.assess_article_quality(source_content, source_url="")
    elif url:
        article = _fetch_article_content(url)
        if not article:
            failure = _core_extractors.article_failure(
                url,
                status="extract_failed",
                error_hint="无法提取文章内容，请在浏览器中打开文章后复制全文导入",
            )
            return jsonify({
                "error": "无法提取文章内容，请检查链接或改用 raw_content 模式",
                "hint": failure["error_hint"],
                "status": failure["status"],
                "error_hint": failure["error_hint"],
                "manual_import_recommended": failure["manual_import_recommended"],
                "source_url": failure["source_url"],
                "extraction_method": failure["extraction_method"],
                "source_quality": failure["quality"],
                "import_mode": import_info["mode"],
                "import_recommendation": import_info["recommendation"],
                "import_message": import_info["message"],
            }), 422

        article = _core_extractors.normalize_article_result(article, fallback_url=url)
        original_title = article["title"]
        original_author = article.get("author", "")
        source_content = article["content"]
        extraction_method = article.get("extraction_method", "unknown")
        source_url = article.get("source_url", url)
        source_quality = article.get("quality") or _core_extractors.assess_article_quality(source_content, source_url=source_url)
    else:
        return jsonify({"error": "url 或 raw_content 至少填一个"}), 400

    # ... rest unchanged ...

    # ── Select style prompt ──
    if style == "b2p":
        system_prompt = B2P_SYSTEM_PROMPT
        style_label = "政策解读 (B2P)"
    elif style == "b2b":
        system_prompt = B2B_SYSTEM_PROMPT
        style_label = "行业深度 (B2B)"
    else:
        system_prompt = B2C_SYSTEM_PROMPT
        style_label = "卖家亲测 (B2C)"

    # ── Build rewrite prompt ──
    angle_instruction = ""
    if custom_angle:
        angle_instruction = f"\n\n**特别要求 — 写作角度**: {custom_angle}"

    rewrite_user_prompt = f"""请根据你的写作风格和结构公式，重写以下文章。

要求：
1. 保留原文核心事实、数据和关键信息
2. 完全替换写作风格、结构和语言节奏
3. 去除原文中的营销广告、CTA、联系方式
4. 如果是政策/法规类内容，确保政策编号、日期、金额等精确数据不丢失
5. 严格事实边界：不得编造原文没有的信息、日期、政策编号、金额、比例、案例、平台动作或专家判断
6. 原文没有的信息一律不要写；如果需要衔接，只能用概括性表达，不能新增具体事实
7. 今天是 {_today_cst_label()}

=== 原文标题 ===
{original_title}

=== 原文作者 ===
{original_author or "未知"}

=== 原文内容 ===
{source_content}{angle_instruction}

=== 要求 ===
直接输出改写后的完整文章（Markdown 格式），以标题行开头。"""

    # ── Call LLM ──
    print(f"  [Rewrite] style={style} ({style_label}), source_len={len(source_content)}")
    rewritten_md = llm_chat_text(
        system=system_prompt,
        user=rewrite_user_prompt,
        temperature=0.7,
        max_tokens=4096,
    )

    if not rewritten_md:
        return jsonify({"error": "AI 改写失败，请重试"}), 500

    # ── Convert to WeChat HTML ──
    rewritten_html = _markdown_to_wechat_html(
        rewritten_md,
        title=original_title,
        theme=theme,
        font_size=15,
    )

    return jsonify({
        "original_title": original_title,
        "original_author": original_author,
        "original_char_count": len(source_content),
        "rewritten_markdown": rewritten_md,
        "rewritten_html": rewritten_html,
        "style": style,
        "style_label": style_label,
        "theme": theme,
        "wechat_compatible": True,
        "char_count": len(rewritten_md),
        "source_url": source_url,
        "extraction_method": extraction_method,
        "source_quality": source_quality,
        "import_mode": import_info["mode"],
        "import_recommendation": import_info["recommendation"],
        "import_message": import_info["message"],
    })


# ── Export ────────────────────────────────────────────

@app.route("/api/export-article", methods=["POST"])
def api_export_article():
    """导出文章。Body: {content, format: "md"|"html"|"txt", filename}"""
    data = request.get_json() or {}
    content = data.get("content", "")
    fmt = data.get("format", "md")
    custom_filename = data.get("filename", "")

    if not content:
        return jsonify({"error": "content 必填"}), 400

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if custom_filename:
        safe_name = re.sub(r"[^\w一-鿿\-_]+", "_", custom_filename)[:60]
    else:
        safe_name = f"article_{timestamp}"

    ext = {"md": ".md", "html": ".html", "txt": ".txt"}.get(fmt, ".md")
    filename = f"{safe_name}{ext}"
    filepath = OUTPUT_DIR / "articles" / filename

    if fmt == "html":
        import markdown
        html_body = markdown.markdown(content, extensions=["tables", "fenced_code"])
        html_full = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="utf-8"><title>{safe_name}</title>
<style>body{{max-width:680px;margin:0 auto;padding:20px;font:16px/1.8 -apple-system,sans-serif;color:#333}}
h1{{font-size:1.5em;color:#CE0E19}}h2{{font-size:1.2em}}table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #ddd;padding:8px;text-align:left}}</style>
</head><body>{html_body}</body></html>"""
        filepath.write_text(html_full, encoding="utf-8")
    elif fmt == "txt":
        # Strip markdown
        txt = re.sub(r"[#*>`~\[\]()|_]", "", content)
        txt = re.sub(r"\n{3,}", "\n\n", txt)
        filepath.write_text(txt, encoding="utf-8")
    else:
        filepath.write_text(content, encoding="utf-8")

    return jsonify({
        "format": fmt, "filename": filename,
        "saved_path": str(filepath), "char_count": len(content),
    })


# ── DOCX 导出 ─────────────────────────────────────

@app.route("/api/export-docx", methods=["POST"])
def api_export_docx():
    """导出 DOCX：将 Markdown 文章 + 图片合并为 Word 文档。

    Body: {
        content: markdown_string,
        title: string,
        images: {cover: url, "0": url, ...}  // 可选
    }
    返回 .docx 二进制文件。
    """
    data = request.get_json() or {}
    content = data.get("content", "")
    title = data.get("title", "跨境电商资讯")
    images = data.get("images", {}) or {}

    if not content:
        return jsonify({"error": "content 必填"}), 400

    try:
        return _build_docx(content, title, images)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"DOCX 生成失败: {str(e)}"}), 500


def _build_docx(content, title, images):
    """Core DOCX generation, separated for error handling.

    Images are placed contextually: each body image includes a section_excerpt
    that tells us which paragraph it was generated for.  We find the best-
    matching paragraph in the article and insert the image right after it.
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ── 自定义样式 ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.space_after = Pt(8)

    # Heading 1 (文章标题)
    h1_style = doc.styles["Heading 1"]
    h1_style.font.name = "微软雅黑"
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0xCE, 0x0E, 0x19)
    h1_style.paragraph_format.space_before = Pt(6)
    h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2_style = doc.styles["Heading 2"]
    h2_style.font.name = "微软雅黑"
    h2_style.font.size = Pt(15)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3_style = doc.styles["Heading 3"]
    h3_style.font.name = "微软雅黑"
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    h3_style.paragraph_format.space_before = Pt(14)
    h3_style.paragraph_format.space_after = Pt(6)

    # ── 下载图片辅助函数 ──
    def _download_image(url):
        """下载图片返回 BytesIO，失败返回 None。"""
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            resp = urllib.request.urlopen(req, timeout=30)
            return BytesIO(resp.read())
        except Exception as e:
            print(f"  [DOCX] 图片下载失败: {e}")
            return None

    def _add_image(doc, url, width_inches=6.0):
        """下载并嵌入图片，居中显示。"""
        img_data = _download_image(url)
        if not img_data:
            return False
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_data, width=Inches(min(width_inches, 6.0)))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            return True
        except Exception as e:
            print(f"  [DOCX] 图片嵌入失败: {e}")
            return False

    # ── 封面图 ──
    cover_url = images.get("cover", "")
    if cover_url:
        _add_image(doc, cover_url, width_inches=5.8)
        # 封面后加分隔
        doc.add_paragraph("")

    # ── 文章标题 ──
    doc.add_heading(title, level=1)

    # ── 收集 body images 及它们的 section_excerpt ──
    # New format:  images["0"] = {"url": "...", "section_excerpt": "..."}
    # Old format:  images["0"] = "http://..."  (plain string, no excerpt)
    _body_images = []  # list of (url, section_excerpt)
    for k in sorted(
        [k for k in images.keys() if k != "cover"],
        key=lambda k: int(k) if k.isdigit() else 999,
    ):
        v = images.get(k)
        if isinstance(v, dict):
            _body_images.append((v.get("url", ""), v.get("section_excerpt", "")))
        else:
            # backward-compat: plain string URL, no excerpt
            _body_images.append((v, ""))

    print(f"  [DOCX] Body images to place: {len(_body_images)}")

    # ── 两遍解析 ──
    # Pass 1: build a list of "blocks" — each block has text and knows how to
    #         write itself to the document later; we annotate each block's
    #         combined text for matching against section_excerpt.
    # Pass 2: for each block we write it, then check whether any unplaced
    #         image's section_excerpt best-matches this block — if so insert
    #         the image right after.

    class _Block:
        """A parsed content block that can write itself to the doc."""
        def __init__(self, block_type, text, data=None):
            self.block_type = block_type  # "h2"|"h3"|"table"|"ol"|"ul"|"para"|"hr"
            self.text = text             # plain text for matching
            self.data = data or {}       # extra rendering data

        def write(self, doc_obj):
            if self.block_type == "h2":
                doc_obj.add_heading(self.text, level=2)
            elif self.block_type == "h3":
                doc_obj.add_heading(self.text, level=3)
            elif self.block_type == "hr":
                doc_obj.add_paragraph("─" * 40)
            elif self.block_type == "table":
                _add_markdown_table(doc_obj, self.data["lines"])
            elif self.block_type == "ol":
                for item_text in self.data["items"]:
                    p = doc_obj.add_paragraph(style="List Number")
                    _add_formatted_run(p, item_text)
            elif self.block_type == "ul":
                for item_text in self.data["items"]:
                    p = doc_obj.add_paragraph(style="List Bullet")
                    _add_formatted_run(p, item_text)
            elif self.block_type == "para":
                p = doc_obj.add_paragraph()
                _add_formatted_run(p, self.text)

    blocks = []
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 候选标题区块 → 跳过
        if line.strip().startswith("### 候选标题"):
            while i < len(lines) and lines[i].strip():
                i += 1
            i += 1
            continue

        # 分隔线
        if line.strip() in ("---", "***", "___"):
            blocks.append(_Block("hr", ""))
            i += 1
            continue

        # H2
        if line.startswith("## ") or line.startswith("##\t"):
            blocks.append(_Block("h2", line[3:].strip()))
            i += 1
            continue

        # H3
        if line.startswith("### ") or line.startswith("###\t"):
            blocks.append(_Block("h3", line[4:].strip()))
            i += 1
            continue

        # 表格
        if line.startswith("|") and line.strip().endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            tbl_text = " ".join(t.strip("| ") for t in table_lines)
            blocks.append(_Block("table", tbl_text, {"lines": table_lines}))
            continue

        # 有序列表
        ol_match = re.match(r"^(\d+)[\.\)]\s+(.+)", line)
        if ol_match:
            items = []
            while i < len(lines):
                m = re.match(r"^(\d+)[\.\)]\s+(.+)", lines[i])
                if not m:
                    break
                items.append(m.group(2))
                i += 1
            blocks.append(_Block("ol", " ".join(items), {"items": items}))
            continue

        # 无序列表
        if line.startswith("- ") or line.startswith("* ") or line.startswith("+ "):
            items = []
            while i < len(lines):
                li = lines[i].strip()
                if li.startswith("- ") or li.startswith("* "):
                    items.append(li[2:])
                elif li.startswith("+ "):
                    items.append(li[2:])
                else:
                    break
                i += 1
            blocks.append(_Block("ul", " ".join(items), {"items": items}))
            continue

        # 普通段落
        para_parts = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", "- ", "* ", "+ ")) and not re.match(r"^\d+[\.\)]", lines[i]):
            if lines[i].strip() in ("---", "***", "___"):
                break
            para_parts.append(lines[i].strip())
            i += 1
        if para_parts:
            blocks.append(_Block("para", " ".join(para_parts)))

    # ── 为每张图片找到最佳匹配 block ──
    def _text_similarity(excerpt, block_text):
        """基于字符 bigram 的文本相似度。

        用 bigram 重叠系数衡量两段文本的关联程度。Bigram 天然惩罚
        过短的文本块（如只有 3-4 个字的标题），避免图片总是匹配到
        标题而非包含更多上下文的正文段落。
        """
        if not excerpt or not block_text:
            return 0.0

        def _bigrams(s):
            return {s[i:i+2] for i in range(len(s) - 1)}

        e_bigrams = _bigrams(excerpt)
        b_bigrams = _bigrams(block_text)
        if not e_bigrams or not b_bigrams:
            return 0.0

        intersection = e_bigrams & b_bigrams
        # Dice-like: 2*|A&B| / (|A| + |B|)
        return 2.0 * len(intersection) / (len(e_bigrams) + len(b_bigrams))

    # assigned_img_idx[block_index] = list of (url,) tuples
    assigned_img_idx = {bi: [] for bi in range(len(blocks))}
    unplaced = []  # images we couldn't match → append at end

    for img_idx, (img_url, excerpt) in enumerate(_body_images):
        if not img_url:
            continue
        if not excerpt:
            # No excerpt → fall back to even distribution
            unplaced.append(img_url)
            continue

        best_block = -1
        best_score = 0.0
        for bi, block in enumerate(blocks):
            score = _text_similarity(excerpt, block.text)
            if score > best_score and score > 0.05:
                best_score = score
                best_block = bi

        if best_block >= 0:
            assigned_img_idx[best_block].append(img_url)
            print(f"  [DOCX] Image {img_idx} → block {best_block} "
                  f"(score={best_score:.3f}, excerpt={excerpt[:40]}...)")
        else:
            unplaced.append(img_url)
            print(f"  [DOCX] Image {img_idx} unmatched, appending at end")

    # ── Pass 2: write blocks, inserting images after matched blocks ──
    for bi, block in enumerate(blocks):
        block.write(doc)
        for img_url in assigned_img_idx.get(bi, []):
            _add_image(doc, img_url, width_inches=5.5)

    # ── Append unplaced images at the end ──
    for img_url in unplaced:
        doc.add_paragraph("")  # spacer
        _add_image(doc, img_url, width_inches=5.5)

    # ── 保存到 BytesIO ──
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r"[^\w一-鿿\-_]+", "_", title)[:40]
    filename = f"{safe_title}_{timestamp}.docx"

    # 同时保存到 output
    save_path = OUTPUT_DIR / "articles" / filename
    save_path.write_bytes(buf.getvalue())

    # 返回文件
    return send_file(
        buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True, download_name=filename,
    )


def _add_formatted_run(paragraph, text):
    """向段落添加带格式的文本 run，支持 **粗体** 标记。"""
    import re as _re_local

    parts = _re_local.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            if inner:
                run = paragraph.add_run(inner)
                run.bold = True
        else:
            run = paragraph.add_run(part)


def _add_markdown_table(doc, table_lines):
    """将 Markdown 表格行转换为 Word 表格。"""
    if len(table_lines) < 2:
        return

    def _parse_row(line):
        cells = [c.strip() for c in line.strip("|").split("|")]
        return cells

    header = _parse_row(table_lines[0])
    # Skip separator line (|---|---|)
    data_rows = []
    for tl in table_lines[1:]:
        if re.match(r"^[\|\s\-:]+$", tl):
            continue
        data_rows.append(_parse_row(tl))

    if not header:
        return

    ncols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols, style="Light Grid Accent 1")
    table.autofit = True

    # Header
    for j, cell_text in enumerate(header):
        if j < ncols:
            cell = table.rows[0].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)

    # Data rows
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

    doc.add_paragraph("")  # spacer after table


# ── Auto Pipeline ─────────────────────────────────────

@app.route("/api/auto-pipeline", methods=["POST"])
def api_auto_pipeline():
    """一键全流程：搜索→生成文章→生成图片。
    Body: {keyword, style: "b2b"|"b2c", generate_images: true}
    """
    data = request.get_json() or {}
    keyword = data.get("keyword", "").strip()
    style = data.get("style", "b2b")
    generate_images = data.get("generate_images", True)

    if not keyword:
        return jsonify({"error": "keyword 必填"}), 400

    steps = []

    # Step 1: Search
    print(f"\n  [Pipeline] Step 1: 搜索 '{keyword}'")
    # Simulate a search request
    with app.test_request_context("/api/search", method="POST", json={"keyword": keyword, "max_results": 10}):
        search_resp = api_search()
    search_data = search_resp.get_json()
    results = search_data.get("results", [])
    steps.append({"step": 1, "name": "搜索热点", "count": len(results), "status": "done"})

    if not results:
        return jsonify({"error": "未找到相关新闻", "steps": steps}), 404

    # Pick the best scored result
    best = results[0]
    steps.append({"step": 2, "name": "选中新闻", "news": best.get("title", ""), "status": "done"})

    # Enrich with topics
    enriched = _enrich_news_with_topics([best], limit=1)
    if enriched:
        best = enriched[0]

    # Step 2: Generate article
    print(f"  [Pipeline] Step 2: 生成 {style.upper()} 文章")
    with app.test_request_context("/api/generate-article", method="POST",
                                  json={"news_item": best, "style": style}):
        article_resp = api_generate_article()
    article_data = article_resp.get_json()
    article_content = article_data.get("content", "")
    steps.append({"step": 3, "name": "生成文章", "chars": len(article_content), "status": "done"})

    # Step 3: Generate images
    image_results = []
    if generate_images and article_content:
        print(f"  [Pipeline] Step 3: 生成配图")
        with app.test_request_context("/api/generate-cover", method="POST",
                                      json={"article_content": article_content}):
            cover_resp = api_generate_cover()
        cover_data = cover_resp.get_json()
        image_results.append({"type": "cover", "data": cover_data})
        steps.append({"step": 4, "name": "生成封面图", "urls": cover_data.get("image_urls", []), "status": "done"})

        with app.test_request_context("/api/generate-image", method="POST",
                                      json={"article_content": article_content, "count": 3}):
            body_resp = api_generate_image()
        body_data = body_resp.get_json()
        image_results.append({"type": "body", "data": body_data})
        steps.append({"step": 5, "name": "生成配图", "count": body_data.get("total", 0), "status": "done"})
    else:
        steps.append({"step": 4, "name": "跳过配图", "status": "skipped"})

    return jsonify({
        "keyword": keyword, "style": style,
        "article": article_data, "images": image_results,
        "steps": steps, "status": "completed",
    })


# ╔══════════════════════════════════════════════════════╗
# ║  Main                                               ║
# ╚══════════════════════════════════════════════════════╝

if __name__ == "__main__":
    import sys, io
    # Fix Windows console encoding for emoji
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

    port = PORT  # from .env or default 8888
    for arg in sys.argv:
        if arg.startswith("--port="):
            port = int(arg.split("=")[1])
    print("=" * 60)
    print("  >> 跨境电商热点内容创作工坊")
    print(f"  http://127.0.0.1:{port}")
    print(f"  DeepSeek: {'OK' if DEEPSEEK_API_KEY else 'MISSING'}")
    print(f"  Ark/Seedream: {'OK' if ARK_API_KEY else 'MISSING'}")
    print("=" * 60)
    app.run(host="0.0.0.0", port=port, debug=False)
